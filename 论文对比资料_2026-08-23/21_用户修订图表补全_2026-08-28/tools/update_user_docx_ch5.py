from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image


VERSION_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = VERSION_DIR / "output" / "行为树论文_中文_用户修订图表补全版_2026-08-28.docx"
FIGURE_56 = VERSION_DIR / "figures" / "figure_5_6_size_all_metrics.png"
FIGURE_57 = VERSION_DIR / "figures" / "figure_5_7_screen_all_metrics.png"


def set_cell_text(cell, text: str, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    if align is not None:
        paragraph.alignment = align
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                if row_index == 0:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(8.5)
                else:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)


def find_table_after_caption(doc: Document, caption_prefix: str):
    found_caption = False
    body = doc.element.body
    table_index = 0
    for child in body.iterchildren():
        if child.tag == qn("w:tbl"):
            if found_caption:
                return doc.tables[table_index]
            table_index += 1
        elif child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.xpath(".//w:t")).strip()
            if text.startswith(caption_prefix):
                found_caption = True
    raise ValueError(f"Could not find table after caption {caption_prefix!r}")


def replace_table_rows(table, rows: list[list[str]]) -> None:
    # Match the table to the required rectangular shape.
    while len(table.columns) < len(rows[0]):
        table.add_column(Pt(60))
    while len(table.columns) > len(rows[0]):
        # Removing a column is rare here; keep extra columns blank rather than risk corrupting layout.
        break
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 or col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.cell(row_idx, col_idx), value, align)
    format_table(table)


def previous_image_paragraph(doc: Document, caption_text_prefix: str):
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().startswith(caption_text_prefix):
            cursor = idx - 1
            while cursor >= 0:
                if paragraphs[cursor]._p.xpath(".//w:drawing"):
                    return paragraphs[cursor]
                cursor -= 1
    raise ValueError(f"Could not find image before caption {caption_text_prefix!r}")


def replace_image_before_caption(doc: Document, caption_prefix: str, image_path: Path) -> None:
    paragraph = previous_image_paragraph(doc, caption_prefix)
    blips = paragraph._p.xpath(".//a:blip")
    if not blips:
        raise ValueError(f"No image relationship found before {caption_prefix!r}")
    rid = blips[0].get(qn("r:embed"))
    image_part = doc.part.related_parts[rid]
    image_part._blob = image_path.read_bytes()

    with Image.open(image_path) as image:
        width_px, height_px = image.size
    cx = 4_500_000
    cy = int(cx * height_px / width_px)
    for extent in paragraph._p.xpath(".//wp:extent"):
        extent.set("cx", str(cx))
        extent.set("cy", str(cy))
    for extent in paragraph._p.xpath(".//a:ext"):
        extent.set("cx", str(cx))
        extent.set("cy", str(cy))


def replace_paragraph_text(doc: Document, old_prefix: str, new_text: str, style: str | None = None) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(old_prefix):
            paragraph.text = new_text
            if style:
                paragraph.style = style
            return
    raise ValueError(f"Paragraph starting {old_prefix!r} not found")


def insert_paragraph_after(paragraph, text: str, style: str = "Normal"):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def add_table_after(doc: Document, paragraph, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if r == 0 or c == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.cell(r, c), value, align)
    format_table(table)
    paragraph._p.addnext(table._tbl)
    return table


def remove_blocks_between(doc: Document, start_text: str, stop_prefixes: tuple[str, ...]) -> object:
    start_para = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == start_text:
            start_para = paragraph
            break
    if start_para is None:
        raise ValueError(f"Could not find {start_text!r}")

    body_children = list(doc.element.body.iterchildren())
    start_index = body_children.index(start_para._p)
    remove = []
    for child in body_children[start_index + 1 :]:
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.xpath(".//w:t")).strip()
            if any(text.startswith(prefix) for prefix in stop_prefixes):
                break
        remove.append(child)
    for child in remove:
        doc.element.body.remove(child)
    return start_para


def main() -> None:
    if not DOCX_PATH.is_file():
        raise FileNotFoundError(DOCX_PATH)
    doc = Document(str(DOCX_PATH))

    replace_table_rows(
        find_table_after_caption(doc, "表 5.2"),
        [
            ["显示优化功能", "主要实验结果", "表明的作用", "主要问题"],
            ["智能拖拽重排", "45组测试中的拖动重叠都被消除；平均移动2.40张周围卡片，平均总移动距离为375.90 px", "能够处理拖动产生的节点遮挡", "会带动少量周围卡片，用户需要接受局部位置变化"],
            ["自适应缩放细节", "卡片面积平均减少45.14%，字段平均减少43.61%", "缩小时可以减少画面占用", "部分文字会暂时隐藏，需要放大后查看细节"],
            ["可读连线覆盖", "平均每个测试画面中有3.62条可见连线受到非端点卡片遮挡；开启后卡片背景调整为72%不透明", "被卡片挡住的连线能够部分显示", "只改善局部连线可见性，不会降低整体拥挤程度"],
            ["相关节点聚焦", "全亮候选节点平均减少58.51%，无关节点被统一淡化", "能够突出选中节点及其相关结构", "不方便同时比较无关分支"],
            ["鱼眼聚焦", "目标节点平均宽度由119.43 px增加到224.55 px", "缩小后仍能临时查看局部节点", "部分测试出现新的局部重叠和层级遮挡"],
        ],
    )

    replace_table_rows(
        find_table_after_caption(doc, "表 5.3"),
        [
            ["资源节点", "Adaptive 面积减少", "Related 全亮候选减少", "Fisheye 宽度增幅", "Smart 平均移动节点", "Smart 平均总移动距离", "Overlay 自然遮挡边比例"],
            ["31", "33.46%", "35.41%", "37.04%", "2.67", "441.15 px", "0.1052"],
            ["61", "33.58%", "57.91%", "54.23%", "1.67", "284.06 px", "0.1869"],
            ["121", "46.07%", "68.74%", "114.94%", "1.67", "284.06 px", "0.1918"],
            ["241", "57.93%", "71.38%", "226.33%", "3.00", "435.11 px", "0.0218"],
            ["364", "54.67%", "59.10%", "300.43%", "3.00", "435.11 px", "0.3677"],
        ],
    )

    replace_table_rows(
        find_table_after_caption(doc, "表 5.4"),
        [
            ["屏幕配置", "Adaptive 面积减少", "Related 全亮候选减少", "Fisheye 宽度增幅", "Smart 平均移动节点", "Smart 平均总移动距离", "Overlay 自然遮挡边比例"],
            ["15.94 in", "57.14%", "54.01%", "258.98%", "2.40", "375.90 px", "0.1682"],
            ["26.96 in", "42.61%", "59.61%", "104.90%", "2.40", "375.90 px", "0.1735"],
            ["31.55 in", "35.67%", "61.90%", "75.90%", "2.40", "375.90 px", "0.1825"],
        ],
    )

    replace_image_before_caption(doc, "图 5.6", FIGURE_56)
    replace_image_before_caption(doc, "图 5.7", FIGURE_57)
    replace_paragraph_text(doc, "图 5.6", "图 5.6　显示优化功能随树规模变化的效果")
    replace_paragraph_text(doc, "图 5.7", "图 5.7　显示优化功能随屏幕尺寸配置变化的效果")
    replace_paragraph_text(
        doc,
        "图5.1比较了节点被拖到同一位置时",
        "图5.1比较了节点被拖到同一位置时，功能开启前后的画面。开启后，系统会移动周围节点来避免重叠。有时为了保持原来的树形结构，同一组中距离较远的节点也会一起移动。因此，智能重排只会限制受影响的范围，但不一定只移动离拖动位置最近的节点。",
        "Normal",
    )

    heading = remove_blocks_between(doc, "5.6　结构化开发者评价", ("6.", "第6章", "第六章"))
    intro = insert_paragraph_after(
        heading,
        "在自动数据之外，本项目还按照固定问题对每项显示优化功能进行了实际使用评价。评价问题包括：该功能主要帮助完成什么操作，开启后是否更容易查看或编辑行为树，使用过程中出现了什么问题，该功能更适合哪种行为树规模，以及该功能更适合哪种屏幕尺寸。表 5.5 记录了对应结果。",
        "Normal",
    )
    caption = insert_paragraph_after(intro, "表 5.5　结构化开发者评价结果", "Table Caption")
    eval_rows = [
        ["显示优化功能", "主要帮助操作", "查看或编辑帮助", "使用中的问题", "适合的树规模", "适合的屏幕尺寸"],
        ["智能拖拽重排", "拖动节点和整理局部分支，避免刚移动后卡片互相遮挡", "对编辑帮助直接，尤其适合需要反复调整节点位置时使用", "有时会带动同组中较远的节点一起移动", "中型到大型树更明显，小树收益较少", "三种屏幕都适合，屏幕尺寸影响不大"],
        ["自适应缩放细节", "缩小时减少卡片面积和字段数量，方便先看整体结构", "对查看全局帮助最大，也能减少缩放后的拥挤感", "缩小时部分参数文字会隐藏，需要放大后再查看", "节点越多越有用，121节点以上更明显", "小屏幕最明显，中屏和大屏也适合作为默认功能"],
        ["可读连线覆盖", "查看穿过卡片背景的父子连线，减少线被卡片完全挡住的情况", "对局部线路判断有帮助，但不是全局整理功能", "如果当前线路没有穿过卡片，效果就不明显", "连线密集或局部遮挡明显的树更适合", "中屏和大屏更容易看到多条线路，小屏主要用于局部查看"],
        ["相关节点聚焦", "选择节点后检查它的父节点、子节点和同级节点", "结构检查时很直观，可以把无关分支压到背景里", "不适合同时比较多个无关分支", "分支越多越有价值，61节点以上开始明显", "中屏和大屏优势更明显，小屏也能减少视觉干扰"],
        ["鱼眼聚焦", "在缩小概览中临时放大鼠标附近的目标节点", "能帮助看清局部节点，但更适合作为临时查看工具", "会带来局部重叠，长时间开启会影响整体结构判断", "大型和超大型树的概览状态更适合", "小屏幕收益最大，大屏幕中更适合按需开启"],
    ]
    table = add_table_after(doc, caption, eval_rows)
    after = insert_paragraph_after(
        caption,
        "总体来看，自适应缩放细节最适合作为一直开启的基础显示优化；智能拖拽重排和相关节点聚焦分别更适合编辑节点位置和检查结构关系；鱼眼聚焦和可读连线覆盖更适合在特定局部问题出现时使用。",
        "Normal",
    )
    # Move the summary paragraph behind the evaluation table.
    table._tbl.addnext(after._p)

    doc.save(str(DOCX_PATH))
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
