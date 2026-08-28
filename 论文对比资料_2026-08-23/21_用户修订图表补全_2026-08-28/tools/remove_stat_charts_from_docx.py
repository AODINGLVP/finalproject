from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


VERSION_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = VERSION_DIR / "output" / "行为树论文_中文_用户修订图表补全版_2026-08-28.docx"


def remove_image_and_caption(doc: Document, caption_prefix: str) -> None:
    caption = None
    caption_index = -1
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(caption_prefix):
            caption = paragraph
            caption_index = index
            break
    if caption is None:
        return

    image = None
    for paragraph in reversed(doc.paragraphs[:caption_index]):
        if paragraph._p.xpath(".//w:drawing"):
            image = paragraph
            break
        if paragraph.text.strip():
            continue
    if image is not None:
        image._p.getparent().remove(image._p)
    caption._p.getparent().remove(caption._p)


def main() -> None:
    doc = Document(str(DOCX_PATH))
    remove_image_and_caption(doc, "图 5.6")
    remove_image_and_caption(doc, "图 5.7")
    doc.save(str(DOCX_PATH))
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
