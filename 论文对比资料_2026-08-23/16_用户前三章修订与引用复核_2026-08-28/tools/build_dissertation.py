#!/usr/bin/env python3
"""Build the current dissertation from tracked Markdown sources.

The Markdown files are the only prose sources.  This script renders an editable
Chinese DOCX and Warwick-template LaTeX/PDF outputs without reading text back
from a PDF or an older DOCX.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SCRIPT = Path(__file__).resolve()
VERSION_DIR = SCRIPT.parents[1]
ROOT = SCRIPT.parents[3]
TEMPLATE_ROOT = ROOT / "thesis_draft"
STYLE_SOURCE = TEMPLATE_ROOT / "chinese" / "uwthesis.sty"
CREST_SOURCE = TEMPLATE_ROOT / "chinese" / "crests" / "officialmono.png"
CREST_DIRECTORY = TEMPLATE_ROOT / "chinese" / "crests"
WORD_FIELD_UPDATER = TEMPLATE_ROOT / "tools" / "update_docx_fields_with_word.ps1"
DEFAULT_TECTONIC = ROOT / ".codex_tmp" / "tectonic-0.17.0" / "tectonic.exe"

DEFAULT_ZH_SOURCE = VERSION_DIR / "source" / "论文内容_中文.md"
DEFAULT_EN_SOURCE = VERSION_DIR / "source" / "dissertation_content_en.md"
DEFAULT_OUTPUT_DIR = VERSION_DIR / "output"

ZH_DOCX_NAME = "行为树论文_中文_用户前三章修订_引用复核版_2026-08-28.docx"
ZH_PDF_NAME = "行为树论文_中文_用户前三章修订_引用复核版_2026-08-28.pdf"
EN_PDF_NAME = "Behaviour_Tree_Dissertation_English_User_Revision_2026-08-28.pdf"

REQUIRED_METADATA = ("title", "author", "student_number", "degree", "department", "submission")
FRONT_HEADINGS_ZH = {"致谢", "声明", "摘要", "缩略语", "缩略词"}
FRONT_HEADINGS_EN = {"acknowledgements", "acknowledgments", "declaration", "abstract", "abbreviations"}
REFERENCE_HEADINGS_ZH = {"参考文献", "文献"}
REFERENCE_HEADINGS_EN = {"references", "bibliography"}

HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
LIST_RE = re.compile(r"^\s*([-+*])\s+(.+)$")
ORDERED_LIST_RE = re.compile(r"^\s*(\d+)[.)]\s+(.+)$")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$")
IMAGE_RE = re.compile(r'^!\[(.*?)\]\((.*?)(?:\s+["\'](.*?)["\'])?\)\s*$')
CAPTION_RE = re.compile(
    r"^(?P<label>图|表|Figure|Table)\s*(?P<number>[A-Z]?\d+(?:\.\d+)*)?\s*[：:\-–—　 ]*\s*(?P<text>.*)$",
    re.IGNORECASE,
)
CHAPTER_ZH_RE = re.compile(r"^第\s*(\d+)\s*章[\s　:：.-]*(.+)$")
CHAPTER_EN_RE = re.compile(r"^(?:Chapter\s+)?(\d+)\s*[.:：\-–— ]+(.+)$", re.IGNORECASE)
EXPLICIT_SECTION_RE = re.compile(r"^(?:[A-Z]\.)?\d+(?:\.\d+){1,4}\s*[　 ]+(.+)$")
INLINE_TOKEN_RE = re.compile(
    r"(\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


@dataclass
class Heading:
    level: int
    text: str


@dataclass
class Paragraph:
    text: str
    quote: bool = False


@dataclass
class ListBlock:
    items: list[str]
    ordered: bool


@dataclass
class TableBlock:
    rows: list[list[str]]
    caption: str = ""


@dataclass
class ImageBlock:
    source: str
    caption: str
    title: str = ""


@dataclass
class CodeBlock:
    text: str
    language: str = ""


Block = Heading | Paragraph | ListBlock | TableBlock | ImageBlock | CodeBlock


@dataclass
class MarkdownDocument:
    source: Path
    metadata: dict[str, str]
    blocks: list[Block] = field(default_factory=list)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def parse_simple_yaml(lines: Sequence[str]) -> dict[str, str]:
    """Parse the flat key/value metadata used by the dissertation sources."""
    result: dict[str, str] = {}
    for raw in lines:
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        if ":" not in line:
            raise ValueError(f"Unsupported metadata line: {raw!r}")
        key, value = line.split(":", 1)
        value = value.strip()
        if len(value) >= 2 and value[0] == value[-1] and value[0] in {'"', "'"}:
            value = value[1:-1]
        result[key.strip()] = value
    return result


def split_table_row(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for char in value:
        if escaped:
            current.append(char)
            escaped = False
        elif char == "\\":
            escaped = True
        elif char == "|":
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(char)
    cells.append("".join(current).strip())
    return cells


def is_table_start(lines: Sequence[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and "|" in lines[index]
        and bool(TABLE_SEPARATOR_RE.match(lines[index + 1]))
    )


def starts_new_block(lines: Sequence[str], index: int) -> bool:
    line = lines[index]
    stripped = line.strip()
    if not stripped:
        return True
    return bool(
        HEADING_RE.match(line)
        or LIST_RE.match(line)
        or ORDERED_LIST_RE.match(line)
        or IMAGE_RE.match(stripped)
        or stripped.startswith("```")
        or is_table_start(lines, index)
    )


def parse_markdown(path: Path) -> MarkdownDocument:
    path = path.resolve()
    if not path.is_file():
        raise FileNotFoundError(path)
    lines = path.read_text(encoding="utf-8-sig").replace("\r\n", "\n").split("\n")
    metadata: dict[str, str] = {}
    index = 0
    if lines and lines[0].strip() == "---":
        try:
            end = next(i for i in range(1, len(lines)) if lines[i].strip() == "---")
        except StopIteration as error:
            raise ValueError(f"Unclosed metadata block in {path}") from error
        metadata = parse_simple_yaml(lines[1:end])
        index = end + 1

    missing = [key for key in REQUIRED_METADATA if not metadata.get(key)]
    if missing:
        raise ValueError(f"{path.name} is missing metadata: {', '.join(missing)}")

    blocks: list[Block] = []
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped or stripped == "---":
            index += 1
            continue

        heading_match = HEADING_RE.match(raw)
        if heading_match:
            blocks.append(Heading(len(heading_match.group(1)), heading_match.group(2).strip()))
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index >= len(lines):
                raise ValueError(f"Unclosed code block in {path}")
            index += 1
            blocks.append(CodeBlock("\n".join(code_lines), language))
            continue

        if is_table_start(lines, index):
            rows = [split_table_row(lines[index])]
            expected_columns = len(rows[0])
            index += 2
            while index < len(lines) and lines[index].strip() and "|" in lines[index]:
                row = split_table_row(lines[index])
                if len(row) != expected_columns:
                    raise ValueError(
                        f"Table row has {len(row)} cells; expected {expected_columns}: {lines[index]!r}"
                    )
                rows.append(row)
                index += 1
            while index < len(lines) and not lines[index].strip():
                index += 1
            caption = ""
            if index < len(lines):
                possible_caption = lines[index].strip()
                match = CAPTION_RE.match(possible_caption)
                if match and match.group("label").lower() in {"表", "table"}:
                    caption = possible_caption
                    index += 1
            blocks.append(TableBlock(rows, caption))
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            alt, source, title = image_match.groups()
            index += 1
            # Some authoring views keep a plain caption immediately after the
            # Markdown image for copy/paste convenience.  The alt text is the
            # canonical caption; suppress only an exactly equivalent duplicate.
            next_content = index
            while next_content < len(lines) and not lines[next_content].strip():
                next_content += 1
            if next_content < len(lines):
                candidate = lines[next_content].strip()
                match = CAPTION_RE.match(candidate)
                if (
                    match
                    and match.group("label").lower() in {"图", "figure"}
                    and captions_equivalent(alt.strip(), candidate, "figure")
                ):
                    index = next_content + 1
            blocks.append(ImageBlock(source.strip(), alt.strip(), (title or "").strip()))
            continue

        list_match = LIST_RE.match(raw)
        ordered_match = ORDERED_LIST_RE.match(raw)
        if list_match or ordered_match:
            ordered = ordered_match is not None
            items: list[str] = []
            while index < len(lines):
                match = ORDERED_LIST_RE.match(lines[index]) if ordered else LIST_RE.match(lines[index])
                if not match:
                    break
                items.append(match.group(2).strip())
                index += 1
            blocks.append(ListBlock(items, ordered))
            continue

        quote = stripped.startswith(">")
        paragraph_lines: list[str] = []
        while index < len(lines) and not starts_new_block(lines, index):
            value = lines[index].strip()
            if quote and value.startswith(">"):
                value = value[1:].lstrip()
            paragraph_lines.append(value)
            index += 1
        if not paragraph_lines:
            paragraph_lines.append(stripped[1:].lstrip() if quote else stripped)
            index += 1
        blocks.append(Paragraph(" ".join(part for part in paragraph_lines if part), quote))

    if not any(isinstance(block, Heading) for block in blocks):
        raise ValueError(f"No Markdown headings found in {path}")
    return MarkdownDocument(path, metadata, blocks)


def strip_explicit_number(text: str) -> str:
    match = EXPLICIT_SECTION_RE.match(text.strip())
    return match.group(1).strip() if match else text.strip()


def chapter_title(text: str, language: str) -> str | None:
    match = CHAPTER_ZH_RE.match(text.strip()) if language == "zh" else CHAPTER_EN_RE.match(text.strip())
    return match.group(2).strip() if match else None


def is_reference_heading(text: str, language: str) -> bool:
    value = text.strip()
    if language == "zh":
        return value in REFERENCE_HEADINGS_ZH
    return value.lower() in REFERENCE_HEADINGS_EN


def is_front_heading(text: str, language: str) -> bool:
    value = text.strip()
    if language == "zh":
        return value in FRONT_HEADINGS_ZH
    return value.lower() in FRONT_HEADINGS_EN


def caption_parts(caption: str, expected: str) -> tuple[str, str]:
    match = CAPTION_RE.match(caption.strip()) if caption else None
    if not match:
        return "", caption.strip()
    label = match.group("label").lower()
    if expected == "figure" and label not in {"图", "figure"}:
        return "", caption.strip()
    if expected == "table" and label not in {"表", "table"}:
        return "", caption.strip()
    return (match.group("number") or "").strip(), (match.group("text") or "").strip()


def captions_equivalent(first: str, second: str, kind: str) -> bool:
    first_number, first_text = caption_parts(first, kind)
    second_number, second_text = caption_parts(second, kind)
    normalize = lambda value: re.sub(r"[\s　:：\-–—]+", "", value).casefold()
    return normalize(first_number) == normalize(second_number) and normalize(first_text) == normalize(second_text)


def resolve_image(document: MarkdownDocument, source: str) -> Path | None:
    if not source:
        return None
    candidate = Path(source.replace("/", os.sep))
    if candidate.is_absolute():
        return candidate if candidate.is_file() else None
    local = (document.source.parent / candidate).resolve()
    if local.is_file():
        return local
    repository = (ROOT / candidate).resolve()
    return repository if repository.is_file() else None


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------


def set_east_asian_font(style_or_run, latin: str, east_asian: str, size: float | None = None) -> None:
    style_or_run.font.name = latin
    if size is not None:
        style_or_run.font.size = Pt(size)
    style_or_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asian)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def add_field(paragraph, instruction: str, placeholder: str = "", hidden: bool = False) -> None:
    begin_run = paragraph.add_run()
    instruction_run = paragraph.add_run()
    separate_run = paragraph.add_run()
    result_run = paragraph.add_run(placeholder)
    end_run = paragraph.add_run()
    for run in (begin_run, instruction_run, separate_run, result_run, end_run):
        if hidden:
            vanish = OxmlElement("w:vanish")
            run._element.get_or_add_rPr().append(vanish)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction_element = OxmlElement("w:instrText")
    instruction_element.set(qn("xml:space"), "preserve")
    instruction_element.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    begin_run._r.append(begin)
    instruction_run._r.append(instruction_element)
    separate_run._r.append(separate)
    end_run._r.append(end)


def add_hyperlink(paragraph, text: str, url: str, *, bold: bool = False, italic: bool = False) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    if bold:
        properties.append(OxmlElement("w:b"))
    if italic:
        properties.append(OxmlElement("w:i"))
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_docx_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(64, 64, 64)
        elif token.startswith("["):
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_page_number_format(section, fmt: str, start: int | None = None) -> None:
    section_pr = section._sectPr
    page_number = section_pr.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        section_pr.append(page_number)
    page_number.set(qn("w:fmt"), fmt)
    if start is not None:
        page_number.set(qn("w:start"), str(start))


def add_footer_page_number(section) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, "PAGE", "1")


def configure_docx(document: Document, language: str) -> None:
    chinese = language == "zh"
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(4)
    section.right_margin = Cm(4)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.different_first_page_header_footer = True
    set_page_number_format(section, "lowerRoman", 1)
    add_footer_page_number(section)

    normal = document.styles["Normal"]
    set_east_asian_font(normal, "Times New Roman", "宋体" if chinese else "Times New Roman", 11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Pt(16.5)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    for level, size in ((1, 18), (2, 14), (3, 12), (4, 11)):
        style = document.styles[f"Heading {level}"]
        set_east_asian_font(style, "Arial", "黑体" if chinese else "Arial", size)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(6)

    for name, base, size, bold, alignment in (
        ("Front Matter Heading", "Heading 1", 16, True, WD_ALIGN_PARAGRAPH.CENTER),
        ("Unlisted Front Heading", "Normal", 16, True, WD_ALIGN_PARAGRAPH.CENTER),
        ("Figure Caption", "Caption", 9.5, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Table Caption", "Caption", 9.5, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Code Block", "Normal", 9, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Block Quote", "Normal", 10.5, False, WD_ALIGN_PARAGRAPH.JUSTIFY),
    ):
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles[base]
        set_east_asian_font(style, "Times New Roman", "宋体" if chinese else "Times New Roman", size)
        style.font.bold = bold
        style.paragraph_format.alignment = alignment
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.keep_with_next = name in {"Figure Caption", "Table Caption"}
        if name == "Block Quote":
            style.paragraph_format.left_indent = Cm(0.8)
            style.paragraph_format.right_indent = Cm(0.8)
        if name == "Code Block":
            style.font.name = "Consolas"
            style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")

    set_update_fields(document)


def add_title_page(document: Document, model: MarkdownDocument, language: str) -> None:
    metadata = model.metadata
    if CREST_SOURCE.is_file():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(CREST_SOURCE), width=Cm(3.2))
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(28)
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run(metadata["title"])
    run.bold = True
    run.font.size = Pt(24)
    set_east_asian_font(run, "Arial", "黑体" if language == "zh" else "Arial")

    by_label = "作者" if language == "zh" else "by"
    degree_label = "提交以申请" if language == "zh" else "Submitted for the degree of"
    for value, size, bold in (
        (f"{by_label} {metadata['author']}", 13, False),
        (f"{degree_label} {metadata['degree']}", 12, False),
        (metadata.get("institution", "University of Warwick"), 12, True),
        (metadata["department"], 11, False),
        (f"{('学号' if language == 'zh' else 'Student number')}: {metadata['student_number']}", 11, False),
        (metadata["submission"], 11, False),
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(9)
        run = paragraph.add_run(value)
        run.font.size = Pt(size)
        run.bold = bold
    document.add_page_break()


def add_docx_field_page(document: Document, heading: str, instruction: str, placeholder: str) -> None:
    # Contents/list headings use the same appearance as other front matter,
    # but must not add themselves to the Table of Contents.
    paragraph = document.add_paragraph(heading, style="Unlisted Front Heading")
    paragraph.paragraph_format.space_after = Pt(12)
    field_paragraph = document.add_paragraph()
    field_paragraph.paragraph_format.first_line_indent = Pt(0)
    add_field(field_paragraph, instruction, placeholder)
    document.add_page_break()


def add_docx_caption(document: Document, caption: str, kind: str, language: str) -> None:
    number, text = caption_parts(caption, kind)
    label = ("图" if kind == "figure" else "表") if language == "zh" else kind.title()
    style = "Figure Caption" if kind == "figure" else "Table Caption"
    paragraph = document.add_paragraph(style=style)
    sequence_name = "Figure" if kind == "figure" else "Table"
    if number:
        add_field(paragraph, f"SEQ {sequence_name} \\* ARABIC", "", hidden=True)
        paragraph.add_run(f"{label} {number}　{text}")
    else:
        paragraph.add_run(f"{label} ")
        add_field(paragraph, f"SEQ {sequence_name} \\* ARABIC", "1")
        if text:
            paragraph.add_run(f"　{text}")


def add_docx_table(document: Document, block: TableBlock, language: str) -> None:
    caption = block.caption or ("表　未命名表格" if language == "zh" else "Table  Untitled table")
    add_docx_caption(document, caption, "table", language)
    table = document.add_table(rows=len(block.rows), cols=len(block.rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(block.rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, "D9E2F3")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_docx_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                run.bold = row_index == 0
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_docx_image(document: Document, model: MarkdownDocument, block: ImageBlock, language: str) -> None:
    image_path = resolve_image(model, block.source)
    if image_path:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        try:
            paragraph.add_run().add_picture(str(image_path), width=Cm(12.5))
        except Exception as error:
            raise RuntimeError(f"Could not insert image {image_path}: {error}") from error
    else:
        placeholder = document.add_table(rows=1, cols=1)
        placeholder.alignment = WD_TABLE_ALIGNMENT.CENTER
        placeholder.style = "Table Grid"
        cell = placeholder.cell(0, 0)
        cell.width = Cm(12.5)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F2F2F2")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        missing_label = "图像占位" if language == "zh" else "Image placeholder"
        paragraph.add_run(f"[{missing_label}: {block.source or block.caption}]").italic = True
    caption = block.caption or ("图　待补图例" if language == "zh" else "Figure  Caption to be supplied")
    add_docx_caption(document, caption, "figure", language)


def render_docx(model: MarkdownDocument, output: Path, language: str) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_docx(document, language)
    add_title_page(document, model, language)
    add_docx_field_page(
        document,
        "目录" if language == "zh" else "Table of Contents",
        'TOC \\o "1-3" \\h \\z \\u',
        "请在 Word 中更新目录域。" if language == "zh" else "Update this field in Word.",
    )
    add_docx_field_page(
        document,
        "图目录" if language == "zh" else "List of Figures",
        'TOC \\h \\z \\c "Figure"',
        "请在 Word 中更新图目录域。" if language == "zh" else "Update this field in Word.",
    )
    add_docx_field_page(
        document,
        "表目录" if language == "zh" else "List of Tables",
        'TOC \\h \\z \\c "Table"',
        "请在 Word 中更新表目录域。" if language == "zh" else "Update this field in Word.",
    )

    main_started = False
    for block in model.blocks:
        if isinstance(block, Heading):
            current_chapter = chapter_title(block.text, language) if block.level == 1 else None
            if current_chapter and not main_started:
                section = document.add_section(WD_SECTION_START.NEW_PAGE)
                section.page_width = Cm(21)
                section.page_height = Cm(29.7)
                section.left_margin = Cm(4)
                section.right_margin = Cm(4)
                section.top_margin = Cm(3)
                section.bottom_margin = Cm(3)
                set_page_number_format(section, "decimal", 1)
                add_footer_page_number(section)
                main_started = True
            if block.level == 1 and is_front_heading(block.text, language):
                if document.paragraphs and document.paragraphs[-1].text:
                    document.add_page_break()
                document.add_paragraph(block.text, style="Front Matter Heading")
            elif block.level == 1 and is_reference_heading(block.text, language):
                document.add_page_break()
                document.add_paragraph(block.text, style="Heading 1")
            else:
                level = min(block.level, 4)
                document.add_paragraph(block.text, style=f"Heading {level}")
        elif isinstance(block, Paragraph):
            paragraph = document.add_paragraph(style="Block Quote" if block.quote else "Normal")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_docx_inline(paragraph, block.text)
        elif isinstance(block, ListBlock):
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                paragraph = document.add_paragraph(style=style)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                add_docx_inline(paragraph, item)
        elif isinstance(block, TableBlock):
            add_docx_table(document, block, language)
        elif isinstance(block, ImageBlock):
            add_docx_image(document, model, block, language)
        elif isinstance(block, CodeBlock):
            paragraph = document.add_paragraph(style="Code Block")
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.add_run(block.text)

    document.core_properties.title = model.metadata["title"]
    document.core_properties.author = model.metadata["author"]
    document.core_properties.subject = "Generated from tracked Markdown; do not edit this generated file as a source."
    document.save(output)
    validate_docx(output, model)


def validate_docx(path: Path, model: MarkdownDocument) -> None:
    if not path.is_file() or path.stat().st_size < 10_000:
        raise RuntimeError(f"DOCX output is missing or unexpectedly small: {path}")
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for key in ("title", "author", "student_number"):
        if model.metadata[key] not in text:
            raise AssertionError(f"DOCX is missing metadata field {key!r}")
    source_headings = [block.text for block in model.blocks if isinstance(block, Heading)]
    missing = [heading for heading in source_headings if heading not in text]
    if missing:
        raise AssertionError(f"DOCX is missing headings: {missing[:5]}")
    expected_tables = sum(isinstance(block, TableBlock) for block in model.blocks)
    # Missing-image placeholders are one-cell tables, hence only a lower bound is checked.
    if len(document.tables) < expected_tables:
        raise AssertionError(f"DOCX contains {len(document.tables)} tables; expected at least {expected_tables}")


# ---------------------------------------------------------------------------
# LaTeX/PDF renderer
# ---------------------------------------------------------------------------


LATEX_ESCAPES = {
    "\\": r"\textbackslash{}",
    "{": r"\{",
    "}": r"\}",
    "$": r"\$",
    "&": r"\&",
    "#": r"\#",
    "%": r"\%",
    "_": r"\_",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def latex_escape(text: str) -> str:
    return "".join(LATEX_ESCAPES.get(char, char) for char in text)


def latex_inline(text: str) -> str:
    pieces: list[str] = []
    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        pieces.append(latex_escape(text[position : match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            pieces.append(r"\textbf{" + latex_escape(token[2:-2]) + "}")
        elif token.startswith("*"):
            pieces.append(r"\emph{" + latex_escape(token[1:-1]) + "}")
        elif token.startswith("`"):
            pieces.append(r"\texttt{" + latex_escape(token[1:-1]) + "}")
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                pieces.append(r"\href{" + latex_escape(link.group(2)) + "}{" + latex_escape(link.group(1)) + "}")
        position = match.end()
    pieces.append(latex_escape(text[position:]))
    return "".join(pieces)


def safe_image_suffix(path: Path) -> str:
    suffix = path.suffix.lower()
    return suffix if suffix in {".png", ".jpg", ".jpeg", ".pdf"} else ".png"


def prepare_latex_directory(staging: Path) -> None:
    staging.mkdir(parents=True, exist_ok=True)
    (staging / "figures").mkdir(exist_ok=True)
    shutil.copy2(STYLE_SOURCE, staging / "uwthesis.sty")
    # The title-page macro also loads a separate Warwick keyline PDF, so copy
    # the complete template crest directory instead of only the central crest.
    shutil.copytree(CREST_DIRECTORY, staging / "crests", dirs_exist_ok=True)


def table_column_spec(rows: list[list[str]]) -> str:
    columns = len(rows[0])
    max_lengths = [max(len(re.sub(r"[`*_]", "", row[index])) for row in rows) for index in range(columns)]
    weights = [max(8.0, min(float(length), 40.0)) for length in max_lengths]
    total = sum(weights)
    fractions = [0.92 * weight / total for weight in weights]
    return "@{}" + " ".join(f"p{{{fraction:.3f}\\textwidth}}" for fraction in fractions) + "@{}"


def render_latex_table(block: TableBlock, language: str, table_index: int) -> list[str]:
    _, caption_text = caption_parts(block.caption, "table")
    if not caption_text:
        caption_text = "未命名表格" if language == "zh" else "Untitled table"
    spec = table_column_spec(block.rows)
    lines = [r"\begin{center}", r"\small", rf"\begin{{longtable}}{{{spec}}}"]
    lines.append(rf"\caption{{{latex_inline(caption_text)}}}\label{{tab:auto-{table_index}}}\\")
    lines.append(r"\toprule")
    lines.append(" & ".join(r"\textbf{" + latex_inline(cell) + "}" for cell in block.rows[0]) + r" \\")
    lines.append(r"\midrule")
    lines.append(r"\endfirsthead")
    lines.append(rf"\multicolumn{{{len(block.rows[0])}}}{{l}}{{\small\itshape " + ("续表" if language == "zh" else "continued") + r"}\\")
    lines.append(r"\toprule")
    lines.append(" & ".join(r"\textbf{" + latex_inline(cell) + "}" for cell in block.rows[0]) + r" \\")
    lines.append(r"\midrule")
    lines.append(r"\endhead")
    for row in block.rows[1:]:
        lines.append(" & ".join(latex_inline(cell) for cell in row) + r" \\")
    lines.extend([r"\bottomrule", r"\end{longtable}", r"\end{center}", ""])
    return lines


def render_latex_image(
    model: MarkdownDocument,
    block: ImageBlock,
    language: str,
    figure_index: int,
    staging: Path,
) -> list[str]:
    _, caption_text = caption_parts(block.caption, "figure")
    if not caption_text:
        caption_text = "待补图例" if language == "zh" else "Caption to be supplied"
    image_path = resolve_image(model, block.source)
    lines = [r"\begin{figure}[htbp]", r"\centering"]
    if image_path:
        destination = staging / "figures" / f"figure_{figure_index:03d}{safe_image_suffix(image_path)}"
        shutil.copy2(image_path, destination)
        lines.append(rf"\includegraphics[width=0.92\textwidth,height=0.72\textheight,keepaspectratio]{{figures/{destination.name}}}")
    else:
        missing = block.source or block.caption
        label = "图像占位" if language == "zh" else "Image placeholder"
        lines.append(
            r"\fbox{\parbox[c][5cm][c]{0.84\textwidth}{\centering\itshape "
            + latex_escape(f"{label}: {missing}")
            + "}}"
        )
    lines.extend(
        [
            rf"\caption{{{latex_inline(caption_text)}}}",
            rf"\label{{fig:auto-{figure_index}}}",
            r"\end{figure}",
            "",
        ]
    )
    return lines


def latex_preamble(model: MarkdownDocument, language: str) -> str:
    metadata = model.metadata
    language_package = r"\usepackage[UTF8,fontset=windows]{ctex}" if language == "zh" else ""
    localized_names = (
        "\n".join(
            [
                r"\renewcommand{\contentsname}{目录}",
                r"\renewcommand{\listfigurename}{图目录}",
                r"\renewcommand{\listtablename}{表目录}",
                r"\renewcommand{\figurename}{图}",
                r"\renewcommand{\tablename}{表}",
            ]
        )
        if language == "zh"
        else ""
    )
    return f"""\\documentclass[11pt,a4paper,oneside,extrafontsizes]{{memoir}}
\\usepackage[margin=4cm]{{uwthesis}}
{language_package}
\\usepackage{{booktabs}}
\\usepackage{{longtable}}
\\usepackage{{array}}
\\usepackage{{graphicx}}
\\usepackage{{enumitem}}
\\usepackage[hidelinks]{{hyperref}}
\\setsecnumdepth{{subsubsection}}
\\settocdepth{{subsubsection}}
\\sloppy
{localized_names}

\\title{{{latex_escape(metadata['title'])}}}
\\author{{{latex_escape(metadata['author'])}}}
\\date{{{latex_escape(metadata['submission'])}}}
\\qualification{{{latex_escape(metadata['degree'])}}}
\\department{{{latex_escape(metadata['department'])}}}
\\studentnumber{{{latex_escape(metadata['student_number'])}}}

\\begin{{document}}
\\frontmatter
\\titlepage[crests/officialmono]
\\tableofcontents
\\clearforchapter\\listoffigures
\\clearforchapter\\listoftables
"""


def render_latex(model: MarkdownDocument, staging: Path, language: str) -> Path:
    prepare_latex_directory(staging)
    lines = [latex_preamble(model, language)]
    main_started = False
    table_index = 0
    figure_index = 0
    for block in model.blocks:
        if isinstance(block, Heading):
            current_chapter = chapter_title(block.text, language) if block.level == 1 else None
            if current_chapter:
                if not main_started:
                    lines.extend([r"\mainmatter", ""])
                    main_started = True
                lines.extend([rf"\chapter{{{latex_inline(current_chapter)}}}", ""])
            elif block.level == 1 and is_front_heading(block.text, language):
                lines.extend([rf"\frontchapter{{{latex_inline(block.text)}}}", ""])
            elif block.level == 1 and is_reference_heading(block.text, language):
                lines.extend(
                    [
                        rf"\chapter*{{{latex_inline(block.text)}}}",
                        rf"\markboth{{{latex_inline(block.text)}}}{{{latex_inline(block.text)}}}",
                        rf"\addcontentsline{{toc}}{{chapter}}{{{latex_inline(block.text)}}}",
                        "",
                    ]
                )
            elif block.level == 1:
                if not main_started:
                    lines.extend([r"\mainmatter", ""])
                    main_started = True
                lines.extend([rf"\chapter{{{latex_inline(block.text)}}}", ""])
            else:
                command = {2: "section", 3: "subsection", 4: "subsubsection", 5: "paragraph", 6: "subparagraph"}[
                    block.level
                ]
                lines.extend([rf"\{command}{{{latex_inline(strip_explicit_number(block.text))}}}", ""])
        elif isinstance(block, Paragraph):
            if block.quote:
                lines.extend([r"\begin{quote}", latex_inline(block.text), r"\end{quote}", ""])
            else:
                lines.extend([latex_inline(block.text), ""])
        elif isinstance(block, ListBlock):
            environment = "enumerate" if block.ordered else "itemize"
            lines.append(rf"\begin{{{environment}}}[leftmargin=*]")
            lines.extend(r"\item " + latex_inline(item) for item in block.items)
            lines.extend([rf"\end{{{environment}}}", ""])
        elif isinstance(block, TableBlock):
            table_index += 1
            lines.extend(render_latex_table(block, language, table_index))
        elif isinstance(block, ImageBlock):
            figure_index += 1
            lines.extend(render_latex_image(model, block, language, figure_index, staging))
        elif isinstance(block, CodeBlock):
            lines.extend([r"\begin{verbatim}", block.text, r"\end{verbatim}", ""])
    lines.extend([r"\end{document}", ""])
    main = staging / "main.tex"
    main.write_text("\n".join(lines), encoding="utf-8", newline="\n")
    return main


def compile_pdf(main: Path, tectonic: Path, destination: Path) -> None:
    if not tectonic.is_file():
        raise FileNotFoundError(f"Portable Tectonic was not found: {tectonic}")
    command = [str(tectonic), "-X", "compile", main.name, "--keep-logs", "--keep-intermediates"]
    completed = subprocess.run(
        command,
        cwd=main.parent,
        text=True,
        encoding="utf-8",
        errors="replace",
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        check=False,
    )
    if completed.returncode != 0:
        tail = "\n".join(completed.stdout.splitlines()[-80:])
        raise RuntimeError(f"Tectonic failed ({completed.returncode}) for {main}:\n{tail}")
    built = main.with_suffix(".pdf")
    if not built.is_file() or built.stat().st_size < 10_000:
        raise RuntimeError(f"Tectonic did not create a usable PDF: {built}")
    log = main.with_suffix(".log")
    if log.is_file():
        log_text = log.read_text(encoding="utf-8", errors="replace")
        fatal = re.search(r"! LaTeX Error:|Emergency stop|Fatal error|Undefined control sequence", log_text)
        if fatal:
            raise RuntimeError(f"Fatal LaTeX marker in {log}: {fatal.group(0)}")
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(built, destination)
    validate_pdf(destination)


def validate_pdf(path: Path) -> None:
    try:
        from pypdf import PdfReader
    except ImportError:
        return
    reader = PdfReader(path)
    if len(reader.pages) < 1:
        raise AssertionError(f"PDF contains no pages: {path}")
    if not (reader.metadata and reader.pages):
        raise AssertionError(f"PDF structure could not be read: {path}")


def update_word_fields(docx_path: Path) -> None:
    if os.name != "nt":
        raise RuntimeError("Word field updating is only available on Windows")
    command = [
        "powershell",
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(WORD_FIELD_UPDATER),
        str(docx_path),
    ]
    completed = subprocess.run(command, text=True, encoding="utf-8", errors="replace", capture_output=True)
    if completed.returncode != 0:
        raise RuntimeError(f"Microsoft Word could not update fields:\n{completed.stdout}\n{completed.stderr}")


def write_manifest(
    output_dir: Path,
    sources: Iterable[Path],
    outputs: Iterable[Path],
) -> Path:
    def manifest_path(path: Path) -> str:
        resolved = path.resolve()
        try:
            value = resolved.relative_to(ROOT)
        except ValueError:
            value = resolved
        return str(value).replace("\\", "/")

    data = {
        "generator": str(SCRIPT.relative_to(ROOT)).replace("\\", "/"),
        "template": {
            "path": str(STYLE_SOURCE.relative_to(ROOT)).replace("\\", "/"),
            "sha256": sha256(STYLE_SOURCE),
        },
        "sources": [
            {"path": manifest_path(path), "sha256": sha256(path)}
            for path in sources
        ],
        "outputs": [
            {"path": path.name, "bytes": path.stat().st_size, "sha256": sha256(path)}
            for path in outputs
            if path.is_file()
        ],
    }
    manifest = output_dir / "build_manifest.json"
    manifest.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest


def build(
    zh_source: Path,
    en_source: Path | None,
    output_dir: Path,
    tectonic: Path,
    *,
    skip_pdf: bool,
    update_fields: bool,
) -> list[Path]:
    if not STYLE_SOURCE.is_file() or not CREST_SOURCE.is_file():
        raise FileNotFoundError("The Warwick template style or crest is missing from thesis_draft")
    output_dir.mkdir(parents=True, exist_ok=True)
    build_dir = output_dir / ".build"
    build_dir.mkdir(exist_ok=True)

    zh_model = parse_markdown(zh_source)
    zh_docx = output_dir / ZH_DOCX_NAME
    render_docx(zh_model, zh_docx, "zh")
    outputs = [zh_docx]
    if update_fields:
        update_word_fields(zh_docx)
        validate_docx(zh_docx, zh_model)

    zh_main = render_latex(zh_model, build_dir / "zh", "zh")
    if not skip_pdf:
        zh_pdf = output_dir / ZH_PDF_NAME
        compile_pdf(zh_main, tectonic, zh_pdf)
        outputs.append(zh_pdf)

    sources = [zh_source]
    if en_source is not None and en_source.is_file():
        en_model = parse_markdown(en_source)
        sources.append(en_source)
        en_main = render_latex(en_model, build_dir / "en", "en")
        if not skip_pdf:
            en_pdf = output_dir / EN_PDF_NAME
            compile_pdf(en_main, tectonic, en_pdf)
            outputs.append(en_pdf)
    elif en_source is not None:
        print(f"WARNING: English source not found; English PDF skipped: {en_source}", file=sys.stderr)

    manifest = write_manifest(output_dir, sources, outputs)
    outputs.append(manifest)
    return outputs


SELF_TEST_ZH = """---
title: 构建器自测论文
author: 测试作者
student_number: 0000000
degree: 游戏工程理学硕士
institution: University of Warwick
department: WMG
submission: 2026年9月
language: zh
---

# 致谢

这是一段致谢。

# 声明

这是一段声明。

# 摘要

这是一段包含 **粗体**、*斜体*、`代码` 和 [链接](https://doi.org/10.1000/test) 的摘要。

# 第1章　引言

## 1.1　测试范围

- 项目符号一
- 项目符号二

| 条件 | 结果 |
| --- | ---: |
| 关闭 | 1 |
| 开启 | 2 |

表 1.1　自测表格

![图 1.1　有意缺失的图像占位](missing-self-test.png)

# 参考文献

1. Example (2026). Test reference. [DOI](https://doi.org/10.1000/test).
"""


SELF_TEST_EN = """---
title: Dissertation Builder Self-Test
author: Test Author
student_number: 0000000
degree: MSc Games Engineering
institution: University of Warwick
department: WMG
submission: September 2026
language: en
---

# Acknowledgements

Acknowledgement text.

# Declaration

Declaration text.

# Abstract

An abstract with **bold**, *italic*, `code`, and a [link](https://doi.org/10.1000/test).

# Chapter 1 Introduction

## 1.1 Scope

- First item
- Second item

| Condition | Result |
| --- | ---: |
| Off | 1 |
| On | 2 |

Table 1.1  Self-test table

![Figure 1.1  Deliberately missing image placeholder](missing-self-test.png)

# References

1. Example (2026). Test reference. [DOI](https://doi.org/10.1000/test).
"""


def self_test(tectonic: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="dissertation-builder-") as temporary:
        root = Path(temporary)
        source_dir = root / "source"
        source_dir.mkdir()
        zh_source = source_dir / "test_zh.md"
        en_source = source_dir / "test_en.md"
        zh_source.write_text(SELF_TEST_ZH, encoding="utf-8")
        en_source.write_text(SELF_TEST_EN, encoding="utf-8")
        outputs = build(zh_source, en_source, root / "output", tectonic, skip_pdf=False, update_fields=False)
        expected_names = {ZH_DOCX_NAME, ZH_PDF_NAME, EN_PDF_NAME, "build_manifest.json"}
        if {path.name for path in outputs} != expected_names:
            raise AssertionError(f"Unexpected self-test outputs: {[path.name for path in outputs]}")
        print("SELF_TEST=PASS")
        for path in outputs:
            print(f"SELF_TEST_OUTPUT={path.name}:{path.stat().st_size}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--zh-source", type=Path, default=DEFAULT_ZH_SOURCE)
    parser.add_argument("--en-source", type=Path, default=DEFAULT_EN_SOURCE)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("--tectonic", type=Path, default=DEFAULT_TECTONIC)
    parser.add_argument("--skip-pdf", action="store_true", help="Build the Chinese DOCX and LaTeX staging files only")
    parser.add_argument(
        "--update-word-fields",
        action="store_true",
        help="Open the generated DOCX invisibly in Microsoft Word and update TOC/list fields",
    )
    parser.add_argument("--self-test", action="store_true", help="Build temporary Chinese/English fixtures and delete them")
    parser.add_argument(
        "--validate-only",
        action="store_true",
        help="Parse and validate available Markdown sources without creating any output files",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.self_test:
        self_test(args.tectonic.resolve())
        return
    en_source = args.en_source.resolve() if args.en_source else None
    if args.validate_only:
        sources = [("zh", args.zh_source.resolve())]
        if en_source is not None and en_source.is_file():
            sources.append(("en", en_source))
        for language, source in sources:
            model = parse_markdown(source)
            counts = {
                "headings": sum(isinstance(block, Heading) for block in model.blocks),
                "paragraphs": sum(isinstance(block, Paragraph) for block in model.blocks),
                "lists": sum(isinstance(block, ListBlock) for block in model.blocks),
                "tables": sum(isinstance(block, TableBlock) for block in model.blocks),
                "figures": sum(isinstance(block, ImageBlock) for block in model.blocks),
            }
            print(f"VALID_SOURCE={language}:{source}")
            print("VALID_BLOCKS=" + language + ":" + json.dumps(counts, ensure_ascii=False, sort_keys=True))
        print("VALIDATION=PASS")
        return
    outputs = build(
        args.zh_source.resolve(),
        en_source,
        args.output_dir.resolve(),
        args.tectonic.resolve(),
        skip_pdf=args.skip_pdf,
        update_fields=args.update_word_fields,
    )
    print("BUILD=PASS")
    for path in outputs:
        print(f"OUTPUT={path}")


if __name__ == "__main__":
    main()
