#!/usr/bin/env python3
"""Build an editable, text-only DOCX from a recorded Chinese LaTeX snapshot.

This intentionally does not extract text from a PDF.  The selected Git revision
is the text authority; an optional PDF hash check only proves that the requested
PDF is the one stored beside that source snapshot.
"""

from __future__ import annotations

import argparse
import hashlib
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
MASTER_PATH = "thesis_draft/chinese/uwthesis.tex"
FRONT_PATH = "thesis_draft/chinese/frontmatter.tex"
BIBLIOGRAPHY_PATH = "thesis_draft/chinese/bibliography.tex"
CHAPTER_PATHS = [
    f"thesis_draft/chinese/chapters/{number:02d}_{name}.tex"
    for number, name in [
        (1, "introduction"),
        (2, "literature_review"),
        (3, "methodology"),
        (4, "method"),
        (5, "results"),
        (6, "discussion"),
        (7, "conclusion"),
    ]
]
APPENDICES_PATH = "thesis_draft/chinese/chapters/appendices.tex"
REMOVED_ENVIRONMENTS = ("figure", "table", "longtable")


@dataclass(frozen=True)
class Block:
    style: str
    text: str


def git_bytes(revision: str, repo_path: str) -> bytes:
    completed = subprocess.run(
        ["git", "show", f"{revision}:{repo_path}"],
        cwd=ROOT,
        check=True,
        capture_output=True,
    )
    return completed.stdout


def source_text(revision: str, repo_path: str) -> str:
    if revision.upper() == "WORKTREE":
        return (ROOT / repo_path).read_text(encoding="utf-8")
    return git_bytes(revision, repo_path).decode("utf-8")


def strip_comments(text: str) -> str:
    cleaned: list[str] = []
    for line in text.splitlines():
        output: list[str] = []
        for index, character in enumerate(line):
            if character == "%":
                slash_count = 0
                cursor = index - 1
                while cursor >= 0 and line[cursor] == "\\":
                    slash_count += 1
                    cursor -= 1
                if slash_count % 2 == 0:
                    break
            output.append(character)
        cleaned.append("".join(output))
    return "\n".join(cleaned)


def remove_environment(text: str, environment: str) -> str:
    pattern = re.compile(
        rf"\\begin\{{{re.escape(environment)}\}}.*?\\end\{{{re.escape(environment)}\}}",
        re.DOTALL,
    )
    previous = None
    while previous != text:
        previous = text
        text = pattern.sub("\n\n", text)
    return text


def parse_bibliography(source: str) -> tuple[dict[str, str], list[str]]:
    source = strip_comments(source)
    pattern = re.compile(
        r"\\bibitem\[([^]]+)]\{([^}]+)}\s*(.*?)"
        r"(?=\n\\bibitem|\n\\end\{thebibliography})",
        re.DOTALL,
    )
    citations: dict[str, str] = {}
    references: list[str] = []
    for match in pattern.finditer(source):
        citation_label = match.group(1).replace("~", " ").strip()
        key = match.group(2).strip()
        label_match = re.fullmatch(r"(.+?)\(([^()]*)\)", citation_label)
        if not label_match:
            raise ValueError(f"Cannot parse citation label: {citation_label}")
        citations[key] = f"{label_match.group(1).strip()}, {label_match.group(2).strip()}"
        references.append(match.group(3).strip())
    if not references:
        raise ValueError("No bibliography entries were parsed")
    return citations, references


def collect_environment_labels(source: str, chapter_code: str) -> dict[str, str]:
    source = strip_comments(source)
    labels: dict[str, str] = {}
    counters = {"table": 0, "figure": 0}
    pattern = re.compile(
        r"\\begin\{(table|longtable|figure)\}(.*?)\\end\{\1\}",
        re.DOTALL,
    )
    for match in pattern.finditer(source):
        kind = "figure" if match.group(1) == "figure" else "table"
        counters[kind] += 1
        number = f"{chapter_code}.{counters[kind]}"
        for label in re.findall(r"\\label\{([^}]+)}", match.group(0)):
            labels[label] = number
    return labels


def collect_all_labels(chapter_sources: list[str], appendix_source: str) -> dict[str, str]:
    labels: dict[str, str] = {}
    for chapter_number, source in enumerate(chapter_sources, start=1):
        labels.update(collect_environment_labels(source, str(chapter_number)))

    chapter_matches = list(re.finditer(r"\\chapter\{[^}]+}", appendix_source))
    for index, match in enumerate(chapter_matches):
        end = chapter_matches[index + 1].start() if index + 1 < len(chapter_matches) else len(appendix_source)
        code = chr(ord("A") + index)
        labels.update(collect_environment_labels(appendix_source[match.start():end], code))
    return labels


def render_inline(
    text: str,
    citations: dict[str, str],
    labels: dict[str, str],
) -> str:
    def citation_replacement(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).split(",")]
        missing = [key for key in keys if key not in citations]
        if missing:
            raise ValueError(f"Unknown citation keys: {missing}")
        return "(" + "; ".join(citations[key] for key in keys) + ")"

    def reference_replacement(match: re.Match[str]) -> str:
        key = match.group(1)
        if key not in labels:
            raise ValueError(f"Unknown cross-reference label: {key}")
        return labels[key]

    text = re.sub(r"\\(?:parencite|citep)\{([^}]+)}", citation_replacement, text)
    text = re.sub(r"\\ref\{([^}]+)}", reference_replacement, text)
    text = re.sub(r"\\href\{([^{}]*)}\{([^{}]*)}", lambda match: match.group(2), text)

    unwrap_pattern = re.compile(r"\\(?:emph|texttt|textsc|mathrm|path|url)\{([^{}]*)}")
    previous = None
    while previous != text:
        previous = text
        text = unwrap_pattern.sub(lambda match: match.group(1), text)

    replacements = {
        r"\times": "×",
        r"\qquad": "，",
        r"\%": "%",
        r"\_": "_",
        r"\&": "&",
        r"\#": "#",
        r"\{": "{",
        r"\}": "}",
        r"\ ": " ",
        r"\(": "",
        r"\)": "",
        r"\[": "",
        r"\]": "",
    }
    for source, replacement in replacements.items():
        text = text.replace(source, replacement)

    text = re.sub(r"_\{([^{}]+)}", r"_\1", text)
    text = text.replace("~", " ").replace("--", "–")
    text = text.replace("{", "").replace("}", "")
    unknown = re.findall(r"\\[A-Za-z@]+", text)
    if unknown:
        raise ValueError(f"Unrendered LaTeX commands {unknown} in: {text}")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def prepare_body_source(source: str, include_conditional_text: bool = False) -> str:
    source = strip_comments(source)
    if include_conditional_text:
        source = re.sub(
            r"\\ifdefined\\omitnontechnicalfrontmatter\s*\\else\s*(.*?)\\fi",
            lambda match: match.group(1),
            source,
            flags=re.DOTALL,
        )
    for environment in REMOVED_ENVIRONMENTS:
        source = remove_environment(source, environment)
    return source


def parse_structured_source(
    source: str,
    chapter_codes: list[str],
    citations: dict[str, str],
    labels: dict[str, str],
    *,
    frontmatter: bool = False,
) -> list[Block]:
    source = prepare_body_source(source, include_conditional_text=frontmatter)
    blocks: list[Block] = []
    buffer: list[str] = []
    formula: list[str] | None = None
    chapter_index = -1
    current_chapter = ""
    section_number = 0
    subsection_number = 0

    def flush() -> None:
        nonlocal buffer
        if not buffer:
            return
        rendered = render_inline(" ".join(buffer), citations, labels)
        if rendered:
            blocks.append(Block("Normal", rendered))
        buffer = []

    for raw_line in source.splitlines():
        line = raw_line.strip()
        if formula is not None:
            if line == r"\]":
                rendered = render_inline(" ".join(formula), citations, labels)
                if rendered:
                    blocks.append(Block("Formula", rendered))
                formula = None
            else:
                formula.append(line)
            continue
        if line == r"\[":
            flush()
            formula = []
            continue
        if not line:
            flush()
            continue

        front_match = re.fullmatch(r"\\frontchapter\{(.+)}", line)
        chapter_match = re.fullmatch(r"\\chapter\{(.+)}", line)
        section_match = re.fullmatch(r"\\section\{(.+)}", line)
        subsection_match = re.fullmatch(r"\\subsection\{(.+)}", line)
        item_match = re.fullmatch(r"\\item\[([^]]*)]\s*(.*)", line)

        if front_match:
            flush()
            blocks.append(Block("Heading 1", render_inline(front_match.group(1), citations, labels)))
        elif chapter_match:
            flush()
            chapter_index += 1
            if chapter_index >= len(chapter_codes):
                raise ValueError("More chapter commands than supplied chapter codes")
            current_chapter = chapter_codes[chapter_index]
            section_number = 0
            subsection_number = 0
            title = render_inline(chapter_match.group(1), citations, labels)
            prefix = f"第{current_chapter}章" if current_chapter.isdigit() else f"附录 {current_chapter}"
            blocks.append(Block("Heading 1", f"{prefix}　{title}"))
        elif section_match:
            flush()
            section_number += 1
            subsection_number = 0
            title = render_inline(section_match.group(1), citations, labels)
            blocks.append(Block("Heading 2", f"{current_chapter}.{section_number}　{title}"))
        elif subsection_match:
            flush()
            subsection_number += 1
            title = render_inline(subsection_match.group(1), citations, labels)
            blocks.append(
                Block(
                    "Heading 3",
                    f"{current_chapter}.{section_number}.{subsection_number}　{title}",
                )
            )
        elif item_match:
            flush()
            label = render_inline(item_match.group(1), citations, labels)
            value = render_inline(item_match.group(2), citations, labels)
            blocks.append(Block("Normal", f"{label}{value}"))
        elif re.fullmatch(r"\\label\{[^}]+}", line):
            continue
        elif line in {r"\begin{description}", r"\end{description}", r"\small"}:
            continue
        elif frontmatter and (
            line == r"\frontmatter"
            or line.startswith(r"\titlepage")
            or line == r"\tableofcontents"
            or line in {r"\clearforchapter\listoffigures", r"\clearforchapter\listoftables"}
        ):
            continue
        elif line.startswith(r"\begin{") or line.startswith(r"\end{"):
            raise ValueError(f"Unprocessed environment line: {line}")
        else:
            buffer.append(line)

    if formula is not None:
        raise ValueError("Unclosed display formula")
    flush()
    return blocks


def parse_metadata(master: str) -> dict[str, str]:
    metadata: dict[str, str] = {}
    for key in ("title", "author", "date", "qualification", "department", "studentnumber"):
        match = re.search(rf"\\{key}\{{([^}}]*)}}", master)
        if not match:
            raise ValueError(f"Missing \\{key} in master document")
        metadata[key] = match.group(1)
    return metadata


def set_east_asian_font(style, latin_name: str, east_asian_name: str, size: float) -> None:
    style.font.name = latin_name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asian_name)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(4)
    section.right_margin = Cm(4)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

    normal = document.styles["Normal"]
    set_east_asian_font(normal, "Times New Roman", "宋体", 11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Pt(16.5)
    normal.paragraph_format.space_after = Pt(0)

    set_east_asian_font(document.styles["Title"], "Times New Roman", "黑体", 20)
    for level, size in ((1, 16), (2, 14), (3, 12)):
        style = document.styles[f"Heading {level}"]
        set_east_asian_font(style, "Times New Roman", "黑体", size)
        style.paragraph_format.first_line_indent = Pt(0)


def build_blocks(
    revision: str,
) -> tuple[dict[str, str], list[Block], dict[str, str]]:
    master = source_text(revision, MASTER_PATH)
    front = source_text(revision, FRONT_PATH)
    bibliography = source_text(revision, BIBLIOGRAPHY_PATH)
    chapter_sources = [source_text(revision, path) for path in CHAPTER_PATHS]
    appendices = source_text(revision, APPENDICES_PATH)

    citations, raw_references = parse_bibliography(bibliography)
    labels = collect_all_labels(chapter_sources, appendices)
    blocks = parse_structured_source(
        front,
        [],
        citations,
        labels,
        frontmatter=True,
    )
    for chapter_number, chapter_source in enumerate(chapter_sources, start=1):
        blocks.extend(
            parse_structured_source(
                chapter_source,
                [str(chapter_number)],
                citations,
                labels,
            )
        )

    blocks.append(Block("Heading 1", "参考文献"))
    for raw_reference in raw_references:
        blocks.append(Block("Normal", render_inline(raw_reference, citations, labels)))

    appendix_count = len(re.findall(r"\\chapter\{", appendices))
    appendix_codes = [chr(ord("A") + index) for index in range(appendix_count)]
    blocks.extend(
        parse_structured_source(
            appendices,
            appendix_codes,
            citations,
            labels,
        )
    )
    return parse_metadata(master), blocks, labels


def write_docx(output: Path, metadata: dict[str, str], blocks: list[Block], revision: str) -> None:
    document = Document()
    configure_document(document)
    document.core_properties.title = metadata["title"]
    document.core_properties.subject = "中文论文纯文字修改版"
    document.core_properties.author = metadata["author"]
    document.core_properties.comments = f"Text source: Git revision {revision}; no PDF text extraction."

    title = document.add_paragraph(metadata["title"], style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(30)

    title_page_lines = [
        "by",
        metadata["author"],
        "A thesis submitted in partial fulfilment",
        "of the requirements for the degree of",
        metadata["qualification"],
        "University of Warwick",
        metadata["department"],
        metadata["date"],
    ]
    for text in title_page_lines:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
    document.add_page_break()

    for block in blocks:
        style = "Normal" if block.style == "Formula" else block.style
        paragraph = document.add_paragraph(block.text, style=style)
        if block.style.startswith("Heading"):
            paragraph.paragraph_format.page_break_before = block.style == "Heading 1"
        elif block.style == "Formula":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def visible_texts(metadata: dict[str, str], blocks: list[Block]) -> list[str]:
    return [
        metadata["title"],
        "by",
        metadata["author"],
        "A thesis submitted in partial fulfilment",
        "of the requirements for the degree of",
        metadata["qualification"],
        "University of Warwick",
        metadata["department"],
        metadata["date"],
        *[block.text for block in blocks],
    ]


def verify_docx(output: Path, expected_texts: list[str]) -> tuple[int, int]:
    document = Document(output)
    actual_texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    if actual_texts != expected_texts:
        mismatch = next(
            (
                index
                for index, pair in enumerate(zip(actual_texts, expected_texts))
                if pair[0] != pair[1]
            ),
            min(len(actual_texts), len(expected_texts)),
        )
        raise AssertionError(
            f"DOCX text mismatch at paragraph {mismatch}: "
            f"actual={actual_texts[mismatch:mismatch + 1]!r} "
            f"expected={expected_texts[mismatch:mismatch + 1]!r}"
        )
    if document.tables:
        raise AssertionError(f"Text-only DOCX unexpectedly contains {len(document.tables)} tables")
    if document.inline_shapes:
        raise AssertionError(
            f"Text-only DOCX unexpectedly contains {len(document.inline_shapes)} inline shapes"
        )
    heading_count = sum(
        1 for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading")
    )
    return len(actual_texts), heading_count


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--revision", default="WORKTREE")
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--reference-pdf", type=Path)
    parser.add_argument("--reference-pdf-repo-path")
    args = parser.parse_args()

    if args.reference_pdf:
        if not args.reference_pdf_repo_path or args.revision.upper() == "WORKTREE":
            raise ValueError(
                "PDF snapshot verification requires --revision and --reference-pdf-repo-path"
            )
        local_pdf = args.reference_pdf.resolve().read_bytes()
        recorded_pdf = git_bytes(args.revision, args.reference_pdf_repo_path)
        local_hash = hashlib.sha256(local_pdf).hexdigest().upper()
        recorded_hash = hashlib.sha256(recorded_pdf).hexdigest().upper()
        if local_hash != recorded_hash:
            raise AssertionError(
                f"Reference PDF does not match revision {args.revision}: "
                f"local={local_hash} recorded={recorded_hash}"
            )
        print(f"REFERENCE_PDF_SHA256={local_hash}")

    metadata, blocks, labels = build_blocks(args.revision)
    output = args.output if args.output.is_absolute() else ROOT / args.output
    write_docx(output, metadata, blocks, args.revision)
    paragraphs, headings = verify_docx(output, visible_texts(metadata, blocks))

    required_text = {
        "致谢正文": "此处预留给作者感谢导师",
        "声明正文": "本论文作为本人申请游戏工程理学硕士学位",
        "摘要正文": "可视化行为树常用于编写游戏智能体的决策",
        "3.6": "3.6　有效性与可复现性",
        "完整参考文献": "Godot Engine contributors (2026b)",
        "附录证据": "多规模显示汇总：",
    }
    all_text = "\n".join(visible_texts(metadata, blocks))
    for name, fragment in required_text.items():
        if fragment not in all_text:
            raise AssertionError(f"Required {name} text is missing: {fragment}")

    digest = hashlib.sha256(output.read_bytes()).hexdigest().upper()
    print(f"OUTPUT={output}")
    print(f"DOCX_SHA256={digest}")
    print(f"TEXT_PARAGRAPHS={paragraphs}")
    print(f"HEADINGS={headings}")
    print(f"CROSS_REFERENCES={len(labels)}")
    print("TABLES=0")
    print("INLINE_SHAPES=0")
    print("DOCX_TEXT_SYNC=PASS")


if __name__ == "__main__":
    main()
