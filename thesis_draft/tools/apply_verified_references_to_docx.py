#!/usr/bin/env python3
"""Create a citation-verified DOCX without changing the accepted source DOCX.

The input DOCX is copied first.  Only two body passages and the fourteen
bibliography paragraphs are then changed.  Tables, images, styles, headings,
and every other paragraph are preserved.
"""

from __future__ import annotations

import argparse
import hashlib
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE


ACADEMIC_REFERENCES = [
    (
        "Colledanchise, M. and Ögren, P. (2018). Behavior Trees in Robotics and AI: "
        "An Introduction. CRC Press. ",
        "https://doi.org/10.1201/9780429489105",
        ".",
    ),
    (
        "Furnas, G. W. (1986). Generalized fisheye views. In: Proceedings of the "
        "SIGCHI Conference on Human Factors in Computing Systems (CHI ’86), pp. 16–23. ",
        "https://doi.org/10.1145/22627.22342",
        ".",
    ),
    (
        "Lamping, J., Rao, R. and Pirolli, P. (1995). A focus+context technique based "
        "on hyperbolic geometry for visualizing large hierarchies. In: Proceedings of "
        "the SIGCHI Conference on Human Factors in Computing Systems (CHI ’95), "
        "pp. 401–408. ",
        "https://doi.org/10.1145/223904.223956",
        ".",
    ),
    (
        "Cockburn, A., Karlson, A. and Bederson, B. B. (2008). A review of "
        "overview+detail, zooming, and focus+context interfaces. ACM Computing Surveys, "
        "41(1), Article 2, pp. 1–31. ",
        "https://doi.org/10.1145/1456650.1456652",
        ".",
    ),
    (
        "Song, H., Kim, B., Lee, B. and Seo, J. (2010). A comparative evaluation on "
        "tree visualization methods for hierarchical structures with large fan-outs. "
        "In: Proceedings of the SIGCHI Conference on Human Factors in Computing Systems "
        "(CHI 2010), pp. 223–232. ",
        "https://doi.org/10.1145/1753326.1753359",
        ".",
    ),
    (
        "Bae, J. and Watson, B. (2011). Developing and evaluating Quilts for the "
        "depiction of large layered graphs. IEEE Transactions on Visualization and "
        "Computer Graphics, 17(12), pp. 2268–2275. ",
        "https://doi.org/10.1109/TVCG.2011.187",
        ".",
    ),
    (
        "Dogrusoz, U., Karacelik, A., Safarli, I., Balci, H., Dervishi, L. and Siper, "
        "M. C. (2018). Efficient methods and readily customizable libraries for "
        "managing complexity of large networks. PLOS ONE, 13(5), e0197238. ",
        "https://doi.org/10.1371/journal.pone.0197238",
        ".",
    ),
    (
        "Misue, K., Eades, P., Lai, W. and Sugiyama, K. (1995). Layout adjustment and "
        "the mental map. Journal of Visual Languages & Computing, 6(2), pp. 183–210. ",
        "https://doi.org/10.1006/jvlc.1995.1010",
        ".",
    ),
    (
        "Reingold, E. M. and Tilford, J. S. (1981). Tidier drawings of trees. IEEE "
        "Transactions on Software Engineering, SE-7(2), pp. 223–228. ",
        "https://doi.org/10.1109/TSE.1981.234519",
        ".",
    ),
    (
        "Plaisant, C., Grosjean, J. and Bederson, B. B. (2002). SpaceTree: Supporting "
        "exploration in large node link tree, design evolution and empirical evaluation. "
        "In: IEEE Symposium on Information Visualization 2002 (INFOVIS 2002), "
        "pp. 57–64. ",
        "https://doi.org/10.1109/INFVIS.2002.1173148",
        ".",
    ),
    (
        "Shneiderman, B. (1996). The eyes have it: A task by data type taxonomy for "
        "information visualizations. In: Proceedings of the 1996 IEEE Symposium on "
        "Visual Languages, pp. 336–343. ",
        "https://doi.org/10.1109/VL.1996.545307",
        ".",
    ),
]

WEB_REFERENCES = [
    (
        "Godot Engine contributors (n.d.-a). EditorPlugin. Godot Engine 4.6 "
        "documentation. Available at: ",
        "https://docs.godotengine.org/en/4.6/classes/class_editorplugin.html",
        " (Accessed: 24 August 2026).",
    ),
    (
        "Godot Engine contributors (n.d.-b). GraphEdit. Godot Engine 4.6 "
        "documentation. Available at: ",
        "https://docs.godotengine.org/en/4.6/classes/class_graphedit.html",
        " (Accessed: 24 August 2026).",
    ),
    (
        "Epic Games (n.d.). Behavior Tree in Unreal Engine—Overview. Available at: ",
        "https://dev.epicgames.com/documentation/en-us/unreal-engine/behavior-tree-in-unreal-engine---overview",
        " (Accessed: 24 August 2026).",
    ),
]

REFERENCES = ACADEMIC_REFERENCES + WEB_REFERENCES
URLS = [url for _, url, _ in REFERENCES]


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(colour)
    properties.append(underline)
    run.append(properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def replace_once(document: Document, old: str, new: str) -> None:
    matches = [paragraph for paragraph in document.paragraphs if old in paragraph.text]
    if len(matches) != 1:
        raise AssertionError(f"Expected one body match for {old!r}, found {len(matches)}")
    paragraph = matches[0]
    paragraph.text = paragraph.text.replace(old, new)


def bibliography_paragraphs(document: Document):
    paragraphs = document.paragraphs
    starts = [i for i, p in enumerate(paragraphs) if p.text.strip() == "参考文献"]
    if len(starts) != 1:
        raise AssertionError(f"Expected one 参考文献 heading, found {len(starts)}")
    start = starts[0] + 1
    end = next(
        (i for i in range(start, len(paragraphs)) if paragraphs[i].style.name == "Heading 1"),
        len(paragraphs),
    )
    result = paragraphs[start:end]
    if len(result) != len(REFERENCES):
        raise AssertionError(
            f"Expected {len(REFERENCES)} bibliography paragraphs, found {len(result)}"
        )
    return result


def hyperlink_targets(document: Document) -> list[str]:
    targets: list[str] = []
    for relationship in document.part.rels.values():
        if relationship.reltype == RELATIONSHIP_TYPE.HYPERLINK and relationship.is_external:
            targets.append(relationship.target_ref)
    return targets


def verify(input_path: Path, output_path: Path, original_texts: list[str]) -> None:
    result = Document(output_path)
    if len(result.paragraphs) != len(original_texts):
        raise AssertionError("Paragraph count changed")
    if result.tables:
        raise AssertionError(f"Unexpected tables: {len(result.tables)}")
    if result.inline_shapes:
        raise AssertionError(f"Unexpected inline shapes: {len(result.inline_shapes)}")

    targets = hyperlink_targets(result)
    missing = [url for url in URLS if url not in targets]
    unexpected = [url for url in targets if url not in URLS]
    if missing or unexpected or len(targets) != len(URLS):
        raise AssertionError(
            f"Hyperlink mismatch: count={len(targets)}, missing={missing}, unexpected={unexpected}"
        )

    text = "\n".join(paragraph.text for paragraph in result.paragraphs)
    required = [
        "(Epic Games, n.d.)",
        "(Godot Engine contributors, n.d.-a, n.d.-b)",
        "2268–2275",
        "10.1145/22627.22342",
        "10.1109/INFVIS.2002.1173148",
        "10.1109/VL.1996.545307",
    ]
    forbidden = [
        "(Epic Games, 2026)",
        "2268–2277",
        "10.1145/22339.22342",
        "/en/stable/classes/class_graphedit.html",
        "/en/stable/classes/class_editorplugin.html",
    ]
    for fragment in required:
        if fragment not in text:
            raise AssertionError(f"Required fragment missing: {fragment}")
    for fragment in forbidden:
        if fragment in text:
            raise AssertionError(f"Obsolete fragment remains: {fragment}")

    before = Document(input_path)
    before_headings = [p.text for p in before.paragraphs if p.style.name.startswith("Heading")]
    after_headings = [p.text for p in result.paragraphs if p.style.name.startswith("Heading")]
    if before_headings != after_headings:
        raise AssertionError("Heading sequence changed")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()

    input_path = args.input.resolve()
    output_path = args.output.resolve()
    if input_path == output_path:
        raise ValueError("Output must not overwrite the author-edited input DOCX")
    if not input_path.is_file():
        raise FileNotFoundError(input_path)

    source = Document(input_path)
    original_texts = [paragraph.text for paragraph in source.paragraphs]
    original_headings = sum(
        1 for paragraph in source.paragraphs if paragraph.style.name.startswith("Heading")
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(input_path, output_path)
    document = Document(output_path)

    replace_once(
        document,
        "Live Debug 也只出现在 Godot 编辑器中，不会加入最终游戏画面。",
        "Live Debug 也只出现在 Godot 编辑器中，不会加入最终游戏画面。"
        "编辑器界面通过 Godot 的 EditorPlugin 扩展机制集成，节点画布以 GraphEdit "
        "为基础(Godot Engine contributors, n.d.-a, n.d.-b)。",
    )
    replace_once(document, "(Epic Games, 2026)", "(Epic Games, n.d.)")

    for paragraph, (prefix, url, suffix) in zip(
        bibliography_paragraphs(document),
        REFERENCES,
        strict=True,
    ):
        paragraph.clear()
        paragraph.add_run(prefix)
        add_hyperlink(paragraph, url, url)
        paragraph.add_run(suffix)

    document.core_properties.comments = (
        "Citation-verified derivative of the author-edited DOCX; LaTeX/PDF not yet synchronized."
    )
    document.save(output_path)
    verify(input_path, output_path, original_texts)

    digest = hashlib.sha256(output_path.read_bytes()).hexdigest().upper()
    print(f"INPUT={input_path}")
    print(f"OUTPUT={output_path}")
    print(f"DOCX_SHA256={digest}")
    print(f"PARAGRAPHS={len(original_texts)}")
    print(f"HEADINGS={original_headings}")
    print(f"REFERENCES={len(REFERENCES)}")
    print(f"HYPERLINKS={len(URLS)}")
    print("TABLES=0")
    print("INLINE_SHAPES=0")
    print("CITATION_VERIFICATION=PASS")


if __name__ == "__main__":
    main()
