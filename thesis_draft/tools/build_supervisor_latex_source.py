#!/usr/bin/env python3
"""Generate the English supervisor-review LaTeX source from the accepted DOCX.

The accepted Chinese DOCX is the authority for scope.  The tracked English
translation supplies the English wording and LaTeX markup.  This generator
applies the same accepted omissions, restores the verified bibliography, and
extracts the embedded evidence figures from the DOCX so that the resulting PDF
does not silently fall back to an older screenshot set.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
ENGLISH_ROOT = ROOT / "thesis_draft" / "english"
TRANSLATED = ENGLISH_ROOT / "translated_content_only"


CHAPTER_NAMES = (
    "01_introduction.tex",
    "02_literature_review.tex",
    "03_methodology.tex",
    "04_method.tex",
    "05_results.tex",
    "06_discussion.tex",
    "07_conclusion.tex",
)


DOCX_MEDIA = {
    "word/media/image2.png": "07_runtime_failure.png",
    "word/media/image3.png": "01_baseline.png",
    "word/media/image4.png": "02_compact.png",
    "word/media/image5.png": "03_fisheye.png",
    "word/media/image6.png": "05_collapsed.png",
    "word/media/image7.png": "10_complex_tree.png",
    "word/media/image8.png": "08_orthogonal_edges.png",
    "word/media/image9.png": "09_bundled_edges.png",
    "word/media/image10.png": "optimization_ratios.png",
}


VERIFIED_URLS = (
    "https://doi.org/10.1201/9780429489105",
    "https://doi.org/10.1145/22627.22342",
    "https://doi.org/10.1145/223904.223956",
    "https://doi.org/10.1145/1456650.1456652",
    "https://doi.org/10.1145/1753326.1753359",
    "https://doi.org/10.1109/TVCG.2011.187",
    "https://doi.org/10.1371/journal.pone.0197238",
    "https://doi.org/10.1006/jvlc.1995.1010",
    "https://doi.org/10.1109/TSE.1981.234519",
    "https://doi.org/10.1109/INFVIS.2002.1173148",
    "https://doi.org/10.1109/VL.1996.545307",
    "https://docs.godotengine.org/en/4.6/classes/class_editorplugin.html",
    "https://docs.godotengine.org/en/4.6/classes/class_graphedit.html",
    "https://dev.epicgames.com/documentation/en-us/unreal-engine/behavior-tree-in-unreal-engine---overview",
)


MAIN_TEX = r"""\documentclass[11pt,a4paper,oneside,extrafontsizes]{memoir}
\usepackage[margin=4cm]{uwthesis}
\usepackage[round,authoryear]{natbib}
\newcommand{\parencite}[1]{\citep{#1}}
\usepackage{booktabs}
\usepackage{multirow}
\usepackage{graphicx}
\usepackage{siunitx}
\usepackage{longtable}
\usepackage[hidelinks]{hyperref}
\usepackage[nameinlink,noabbrev]{cleveref}
\setlength{\emergencystretch}{2em}

\title{Display Optimisation and Experimental Evaluation of Large Visual Behaviour Trees in Godot 4.6}
\author{[AUTHOR NAME]}
\date{September 2026}
\qualification{Master of Science in Games Engineering}
\department{WMG, University of Warwick}
\studentnumber{[STUDENT NUMBER]}

\hypersetup{
  pdftitle={Display Optimisation and Experimental Evaluation of Large Visual Behaviour Trees in Godot 4.6},
  pdfauthor={[AUTHOR NAME]},
  bookmarksnumbered=true
}

\begin{document}
\input{frontmatter}
\mainmatter
\input{chapters/01_introduction}
\input{chapters/02_literature_review}
\input{chapters/03_methodology}
\input{chapters/04_method}
\input{chapters/05_results}
\input{chapters/06_discussion}
\input{chapters/07_conclusion}
\input{bibliography}
\end{document}
"""


BIBLIOGRAPHY_TEX = r"""\begin{thebibliography}{99}
\bibitem[Colledanchise and Ögren(2018)]{colledanchise2018bt}
Colledanchise, M. and Ögren, P. (2018). \emph{Behavior Trees in Robotics and AI: An Introduction}. CRC Press. \url{https://doi.org/10.1201/9780429489105}.

\bibitem[Furnas(1986)]{furnas1986fisheye}
Furnas, G. W. (1986). Generalized fisheye views. In: \emph{Proceedings of the SIGCHI Conference on Human Factors in Computing Systems (CHI '86)}, pp. 16--23. \url{https://doi.org/10.1145/22627.22342}.

\bibitem[Lamping et~al.(1995)]{lamping1995hyperbolic}
Lamping, J., Rao, R. and Pirolli, P. (1995). A focus+context technique based on hyperbolic geometry for visualizing large hierarchies. In: \emph{Proceedings of the SIGCHI Conference on Human Factors in Computing Systems (CHI '95)}, pp. 401--408. \url{https://doi.org/10.1145/223904.223956}.

\bibitem[Cockburn et~al.(2008)]{cockburn2008review}
Cockburn, A., Karlson, A. and Bederson, B. B. (2008). A review of overview+detail, zooming, and focus+context interfaces. \emph{ACM Computing Surveys}, 41(1), Article 2, pp. 1--31. \url{https://doi.org/10.1145/1456650.1456652}.

\bibitem[Song et~al.(2010)]{song2010largefanout}
Song, H., Kim, B., Lee, B. and Seo, J. (2010). A comparative evaluation on tree visualization methods for hierarchical structures with large fan-outs. In: \emph{Proceedings of the SIGCHI Conference on Human Factors in Computing Systems (CHI 2010)}, pp. 223--232. \url{https://doi.org/10.1145/1753326.1753359}.

\bibitem[Bae and Watson(2011)]{bae2011quilts}
Bae, J. and Watson, B. (2011). Developing and evaluating Quilts for the depiction of large layered graphs. \emph{IEEE Transactions on Visualization and Computer Graphics}, 17(12), pp. 2268--2275. \url{https://doi.org/10.1109/TVCG.2011.187}.

\bibitem[Dogrusoz et~al.(2018)]{dogrusoz2018complexity}
Dogrusoz, U., Karacelik, A., Safarli, I., Balci, H., Dervishi, L. and Siper, M. C. (2018). Efficient methods and readily customizable libraries for managing complexity of large networks. \emph{PLOS ONE}, 13(5), e0197238. \url{https://doi.org/10.1371/journal.pone.0197238}.

\bibitem[Misue et~al.(1995)]{misue1995mentalmap}
Misue, K., Eades, P., Lai, W. and Sugiyama, K. (1995). Layout adjustment and the mental map. \emph{Journal of Visual Languages \& Computing}, 6(2), pp. 183--210. \url{https://doi.org/10.1006/jvlc.1995.1010}.

\bibitem[Reingold and Tilford(1981)]{reingold1981tidier}
Reingold, E. M. and Tilford, J. S. (1981). Tidier drawings of trees. \emph{IEEE Transactions on Software Engineering}, SE-7(2), pp. 223--228. \url{https://doi.org/10.1109/TSE.1981.234519}.

\bibitem[Plaisant et~al.(2002)]{plaisant2002spacetree}
Plaisant, C., Grosjean, J. and Bederson, B. B. (2002). SpaceTree: Supporting exploration in large node link tree, design evolution and empirical evaluation. In: \emph{IEEE Symposium on Information Visualization 2002 (INFOVIS 2002)}, pp. 57--64. \url{https://doi.org/10.1109/INFVIS.2002.1173148}.

\bibitem[Shneiderman(1996)]{shneiderman1996eyes}
Shneiderman, B. (1996). The eyes have it: A task by data type taxonomy for information visualizations. In: \emph{Proceedings of the 1996 IEEE Symposium on Visual Languages}, pp. 336--343. \url{https://doi.org/10.1109/VL.1996.545307}.

\bibitem[Godot Engine contributors(n.d.-a)]{godoteditorplugin}
Godot Engine contributors (n.d.-a). EditorPlugin. Godot Engine 4.6 documentation. Available at: \url{https://docs.godotengine.org/en/4.6/classes/class_editorplugin.html} (Accessed: 24 August 2026).

\bibitem[Godot Engine contributors(n.d.-b)]{godotgraphedit}
Godot Engine contributors (n.d.-b). GraphEdit. Godot Engine 4.6 documentation. Available at: \url{https://docs.godotengine.org/en/4.6/classes/class_graphedit.html} (Accessed: 24 August 2026).

\bibitem[Epic Games(n.d.)]{unrealbt}
Epic Games (n.d.). Behavior Tree in Unreal Engine---Overview. Available at: \url{https://dev.epicgames.com/documentation/en-us/unreal-engine/behavior-tree-in-unreal-engine---overview} (Accessed: 24 August 2026).
\end{thebibliography}
"""


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def replace_once(text: str, old: str, new: str, *, label: str) -> str:
    count = text.count(old)
    if count != 1:
        raise AssertionError(f"Expected one {label}, found {count}")
    return text.replace(old, new, 1)


def truncate_before(text: str, marker: str, *, label: str) -> str:
    count = text.count(marker)
    if count != 1:
        raise AssertionError(f"Expected one {label}, found {count}")
    return text.split(marker, 1)[0].rstrip() + "\n"


def stack_pair(text: str, first_file: str, first_label: str, second_file: str, second_label: str) -> str:
    old = (
        f"\\begin{{minipage}}{{0.48\\textwidth}}\\centering\\includegraphics[width=\\linewidth]"
        f"{{figures/{first_file}}}\\\\[-1mm]\\small {first_label}\\end{{minipage}}\\hfill\n"
        f"\\begin{{minipage}}{{0.48\\textwidth}}\\centering\\includegraphics[width=\\linewidth]"
        f"{{figures/{second_file}}}\\\\[-1mm]\\small {second_label}\\end{{minipage}}"
    )
    new = (
        f"\\includegraphics[width=0.92\\textwidth]{{figures/{first_file}}}\n"
        f"\\par\\smallskip{{\\small {first_label}}}\\par\\medskip\n"
        f"\\includegraphics[width=0.92\\textwidth]{{figures/{second_file}}}\n"
        f"\\par\\smallskip{{\\small {second_label}}}"
    )
    return replace_once(text, old, new, label=f"paired figure block {first_file}/{second_file}")


def transform_chapter(name: str, text: str) -> str:
    if name == "02_literature_review.tex":
        text = truncate_before(
            text,
            r"\section{Research Gap and Design Implications}",
            label="removed Section 2.6",
        )

    if name == "03_methodology.tex":
        paragraph = (
            "The research is divided into five stages. First, the research question was determined from supervisor "
            "feedback and dissertation requirements; second, the behaviour-tree runtime and game demonstration were "
            "implemented; third, editing, saving, undo and live-debugging functions were completed; fourth, display "
            "optimisations that could be enabled and disabled independently were added; and fifth, fixed behaviour "
            "trees and automated scripts were used to compare results before and after optimisation. Tests were rerun "
            "after each new feature was added to confirm that existing functionality had not been broken.\n\n"
        )
        text = replace_once(text, paragraph, "", label="removed five-stage methodology paragraph")
        text = replace_once(
            text,
            " The results can explain only geometric display differences produced by different screen sizes and cannot directly demonstrate that a large screen is easier to use.",
            "",
            label="removed screen-size limitation sentence",
        )

    if name == "04_method.tex":
        text = replace_once(
            text,
            "Live Debug also appears only in the Godot editor and is not added to the final game view.",
            "Live Debug also appears only in the Godot editor and is not added to the final game view. "
            "The editor interface is integrated through Godot's EditorPlugin extension mechanism, and its node canvas "
            "is based on GraphEdit \\parencite{godoteditorplugin,godotgraphedit}.",
            label="Godot editor citations",
        )
        text = replace_once(
            text,
            "{figures/architecture.png}",
            "{figures/architecture_en.png}",
            label="English architecture figure",
        )
        cross_references = (
            ("None of this debugging content appears in the game view.", " An example is shown in Figure~\\ref{fig:livedebug}."),
            ("reducing reliance on a single colour channel.", " The two information-density states are shown in Figure~\\ref{fig:compact}."),
            ("the original node positions and parent--child relationships are still used.", " Fisheye and collapse examples are shown in Figure~\\ref{fig:focuscollapse}."),
            ("Automatic arrangement still preserves the left-to-right execution order of siblings.", " The related navigation controls are shown in Figure~\\ref{fig:navigation}."),
            ("branch actually traversed by the current Tick.", " Connection and runtime-state examples are shown in Figure~\\ref{fig:edgesruntime}."),
        )
        for original, addition in cross_references:
            text = replace_once(text, original, original + addition, label=f"cross-reference after {original[:28]}")
        text = stack_pair(text, "01_baseline.png", "(a) Baseline", "02_compact.png", "(b) Compact cards")
        text = stack_pair(text, "03_fisheye.png", "(a) Fisheye focus and context", "05_collapsed.png", "(b) Subtree collapse and summary")
        text = stack_pair(text, "08_orthogonal_edges.png", "(a) Orthogonal connections and active path", "09_bundled_edges.png", "(b) Edge bundling and failure annotation")
        text = text.replace(
            "Full-detail baseline and compact display of the same behaviour tree.",
            "Full-detail Baseline and Compact Cards for the same behaviour tree.",
        )

    if name == "05_results.tex":
        text = replace_once(
            text,
            "Every condition preserved node positions and left-to-right execution order.",
            "Every condition preserved node positions and left-to-right execution order. "
            "The reductions across scales are shown in Figure~\\ref{fig:optimization}.",
            label="Figure 5.1 cross-reference",
        )
        text = text.replace(
            "Complex-tree coverage on three physical screen sizes. Resolution is audit metadata, not an experimental treatment.",
            "Complex-tree coverage on three physical screen sizes.",
        )
        text = truncate_before(
            text,
            r"\section{Results for the Research Question}",
            label="removed Section 5.4",
        )

    if name == "07_conclusion.tex":
        text = replace_once(
            text,
            "Functionality testing and the playable game confirmed that the plugin could serve as an experimental "
            "platform. The actual subject investigated by this dissertation is display optimisation, not the basic "
            "functionality of behaviour trees themselves.\n\n",
            "",
            label="removed platform-summary paragraph",
        )
        text = replace_once(
            text,
            " Because no human experiment was conducted, this dissertation cannot demonstrate that user understanding improved.",
            "",
            label="removed human-study sentence",
        )
        text = truncate_before(text, r"\section{Contributions}", label="removed Sections 7.2--7.5")

    return text


def external_hyperlinks(document: Document) -> list[str]:
    return [
        relationship.target_ref
        for relationship in document.part.rels.values()
        if relationship.reltype.endswith("/hyperlink") and relationship.is_external
    ]


def validate_source_docx(path: Path) -> dict[str, object]:
    document = Document(path)
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    text = "\n".join(p.text for p in document.paragraphs)
    if len(headings) != 45:
        raise AssertionError(f"Expected 45 accepted Chinese headings, found {len(headings)}")
    if len(document.tables) != 5:
        raise AssertionError(f"Expected 5 accepted data tables, found {len(document.tables)}")
    if len(document.inline_shapes) != 10:
        raise AssertionError(f"Expected 10 accepted figures, found {len(document.inline_shapes)}")
    for forbidden in ("2.6　研究空白与设计启示", "附录 A", "附录 B", "附录 C"):
        if forbidden in text:
            raise AssertionError(f"Accepted DOCX unexpectedly contains removed content: {forbidden}")
    links = external_hyperlinks(document)
    if len(links) != len(VERIFIED_URLS) or set(links) != set(VERIFIED_URLS):
        raise AssertionError("Accepted DOCX bibliography links do not match the verified 14-link set")
    return {
        "sha256": sha256(path),
        "headings": len(headings),
        "tables": len(document.tables),
        "images": len(document.inline_shapes),
        "hyperlinks": len(links),
    }


def validate_latex(output_dir: Path) -> dict[str, object]:
    main = (output_dir / "supervisor_review.tex").read_text(encoding="utf-8")
    chapter_text = "\n".join(
        (output_dir / "chapters" / name).read_text(encoding="utf-8") for name in CHAPTER_NAMES
    )
    bibliography = (output_dir / "bibliography.tex").read_text(encoding="utf-8")
    combined = main + "\n" + chapter_text + "\n" + bibliography
    forbidden = (
        r"\section{Research Gap and Design Implications}",
        r"\section{Results for the Research Question}",
        r"\section{Contributions}",
        r"\section{Limitations}",
        r"\section{Future Work}",
        r"\section{Final Conclusion}",
        r"\appendix",
        "Draft Completion Checklist",
        "Acknowledgements",
        "Declaration",
    )
    for fragment in forbidden:
        if fragment in combined:
            raise AssertionError(f"Supervisor LaTeX retained removed content: {fragment}")

    if chapter_text.count(r"\chapter{") != 7:
        raise AssertionError("Expected seven dissertation chapters")
    if chapter_text.count(r"\begin{table}") != 5:
        raise AssertionError("Expected five data tables")
    if chapter_text.count(r"\begin{figure}") != 7:
        raise AssertionError("Expected seven numbered figures")
    if chapter_text.count(r"\includegraphics") != 10:
        raise AssertionError("Expected ten embedded evidence images")

    citation_keys: set[str] = set()
    for match in re.finditer(r"\\(?:paren)?cite\{([^}]+)\}", chapter_text):
        citation_keys.update(key.strip() for key in match.group(1).split(","))
    bibliography_keys = set(re.findall(r"\\bibitem(?:\[[^]]+\])?\{([^}]+)\}", bibliography))
    if citation_keys - bibliography_keys:
        raise AssertionError(f"Citations without bibliography entries: {sorted(citation_keys - bibliography_keys)}")
    if len(bibliography_keys) != 14:
        raise AssertionError(f"Expected 14 bibliography entries, found {len(bibliography_keys)}")
    for url in VERIFIED_URLS:
        if url not in bibliography:
            raise AssertionError(f"Verified bibliography URL is missing: {url}")
    return {
        "chapters": 7,
        "tables": 5,
        "numbered_figures": 7,
        "embedded_images": 10,
        "bibliography_entries": len(bibliography_keys),
        "citation_keys": len(citation_keys),
    }


def generate(source_docx: Path, output_dir: Path) -> None:
    source_stats = validate_source_docx(source_docx)
    chapters_dir = output_dir / "chapters"
    figures_dir = output_dir / "figures"
    crests_dir = output_dir / "crests"
    for directory in (output_dir, chapters_dir, figures_dir, crests_dir):
        directory.mkdir(parents=True, exist_ok=True)

    shutil.copy2(ENGLISH_ROOT / "uwthesis.sty", output_dir / "uwthesis.sty")
    for crest in (ENGLISH_ROOT / "crests").iterdir():
        if crest.is_file():
            shutil.copy2(crest, crests_dir / crest.name)
    shutil.copy2(
        ENGLISH_ROOT / "figures" / "architecture_en.png",
        figures_dir / "architecture_en.png",
    )

    with ZipFile(source_docx) as archive:
        names = set(archive.namelist())
        missing = sorted(set(DOCX_MEDIA) - names)
        if missing:
            raise AssertionError(f"Accepted DOCX is missing embedded media: {missing}")
        for member, target_name in DOCX_MEDIA.items():
            (figures_dir / target_name).write_bytes(archive.read(member))

    for name in CHAPTER_NAMES:
        source = (TRANSLATED / "chapters" / name).read_text(encoding="utf-8")
        transformed = transform_chapter(name, source)
        (chapters_dir / name).write_text(transformed, encoding="utf-8", newline="\n")

    (output_dir / "frontmatter.tex").write_text(
        (TRANSLATED / "frontmatter.tex").read_text(encoding="utf-8"),
        encoding="utf-8",
        newline="\n",
    )
    (output_dir / "bibliography.tex").write_text(BIBLIOGRAPHY_TEX, encoding="utf-8", newline="\n")
    (output_dir / "supervisor_review.tex").write_text(MAIN_TEX, encoding="utf-8", newline="\n")

    latex_stats = validate_latex(output_dir)
    manifest = {
        "source_docx": str(source_docx.resolve()),
        "source": source_stats,
        "latex": latex_stats,
        "scope": {
            "language": "English",
            "audience": "UK supervisor review",
            "omitted": [
                "Acknowledgements",
                "Declaration",
                "Section 2.6 removed by the author",
                "Section 5.4 removed by the author",
                "Sections 7.2--7.5 removed by the author",
                "All appendices removed by the author",
            ],
        },
    }
    (output_dir / "source_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(f"SUPERVISOR_LATEX_SOURCE={output_dir.resolve()}")
    print(f"SOURCE_DOCX_SHA256={source_stats['sha256']}")
    print("SUPERVISOR_LATEX_VALIDATION=PASS")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-docx", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    args = parser.parse_args()
    source_docx = args.source_docx if args.source_docx.is_absolute() else ROOT / args.source_docx
    output_dir = args.output_dir if args.output_dir.is_absolute() else ROOT / args.output_dir
    if not source_docx.is_file():
        raise FileNotFoundError(source_docx)
    generate(source_docx.resolve(), output_dir.resolve())


if __name__ == "__main__":
    main()
