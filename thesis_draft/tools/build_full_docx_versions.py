#!/usr/bin/env python3
"""Build complete Chinese and supervisor-facing English dissertation DOCX files.

The Chinese file is enriched from the accepted, citation-verified text DOCX.
The English file reuses the tracked English translation, removes author-only
front matter and checklist material, and is filtered to the same chapter scope
as the accepted Chinese DOCX.  Neither path uses PDF text extraction.
"""

from __future__ import annotations

import argparse
import hashlib
import re
import shutil
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TOOLS = Path(__file__).resolve().parent
if str(TOOLS) not in sys.path:
    sys.path.insert(0, str(TOOLS))

import build_text_only_docx as text_builder  # noqa: E402
from apply_verified_references_to_docx import (  # noqa: E402
    REFERENCES,
    URLS,
    add_hyperlink,
)


THESIS_FIGURES = ROOT / "thesis_draft" / "chinese" / "figures"
CURRENT_VISUALS = ROOT / "testgame" / "testgame" / "test_results" / "visual"
ENGLISH_SOURCE = ROOT / "thesis_draft" / "english" / "translated_content_only"
ENGLISH_MASTER = ROOT / "thesis_draft" / "english" / "uwthesis.tex"


FIGURE_PATHS = {
    "architecture.png": THESIS_FIGURES / "architecture.png",
    "optimization_ratios.png": THESIS_FIGURES / "optimization_ratios.png",
    "01_baseline.png": CURRENT_VISUALS / "01_baseline.png",
    "02_compact.png": CURRENT_VISUALS / "02_compact.png",
    "03_fisheye.png": CURRENT_VISUALS / "03_fisheye.png",
    "05_collapsed.png": CURRENT_VISUALS / "05_collapsed.png",
    "07_runtime_failure.png": CURRENT_VISUALS / "07_runtime_failure.png",
    "08_orthogonal_edges.png": CURRENT_VISUALS / "08_orthogonal_edges.png",
    "09_bundled_edges.png": CURRENT_VISUALS / "09_bundled_edges.png",
    "10_complex_tree.png": CURRENT_VISUALS / "10_complex_overview_default.png",
}


CHINESE_PATH_BEFORE = (
    "物理屏幕证据：实验根目录为 research/display_optimization/ "
    "physical_screen_size_experiment/；其中 data/2026-08-23_three_displays/ Session "
    "包含 combined_raw.csv、screen_size_comparison.csv 与 report_zh.md。"
)
CHINESE_PATH_AFTER = (
    "物理屏幕证据：实验根目录为 research/display_optimization/"
    "physical_screen_size_experiment/；data/2026-08-23_three_displays/ 目录包含 "
    "combined_raw.csv、screen_size_comparison.csv 与 report_zh.md。"
)
ENGLISH_PATH_BEFORE = (
    "Physical-screen evidence:The experiment root is research/display_optimization/ "
    "physical_screen_size_experiment/; its data/2026-08-23_three_displays/ Session "
    "contains combined_raw.csv, screen_size_comparison.csv and report_zh.md."
)
ENGLISH_PATH_AFTER = (
    "Physical-screen evidence: the experiment root is research/display_optimization/"
    "physical_screen_size_experiment/; the data/2026-08-23_three_displays/ directory "
    "contains combined_raw.csv, screen_size_comparison.csv and report_zh.md."
)


CROSS_REFERENCES_ZH = (
    ("游戏画面中不会出现这些调试内容", "该显示示例见图4.2。"),
    ("减少只依赖单一颜色通道的问题", "两种信息密度状态见图4.3。"),
    ("展开后仍使用原有节点位置和父子关系", "鱼眼与折叠示例见图4.4。"),
    ("自动排列仍保持同级节点从左到右的执行顺序", "相关导航控件见图4.5。"),
    ("实际经过的分支", "连线与运行状态示例见图4.6。"),
    ("所有条件都保持了节点位置和从左到右的执行顺序", "各规模的降低比例见图5.1。"),
)


CROSS_REFERENCES_EN = (
    ("None of this debugging content appears in the game view", " An example is shown in Figure 4.2."),
    ("reducing reliance on a single colour channel", " The two information-density states are shown in Figure 4.3."),
    ("the original node positions and parent–child relationships are still used", " Fisheye and collapse examples are shown in Figure 4.4."),
    ("Automatic arrangement still preserves the left-to-right execution order of siblings", " The related navigation controls are shown in Figure 4.5."),
    ("branch actually traversed by the current Tick", " Connection and runtime-state examples are shown in Figure 4.6."),
    ("Every condition preserved node positions and left-to-right execution order", " The reductions across scales are shown in Figure 5.1."),
)


TABLES_ZH = {
    "3.1": [
        ["条件", "实验中的具体操作", "界面变化与主要测量目的"],
        ["Baseline", "关闭本实验涉及的显示优化，显示完整卡片和全部分支。", "提供原始显示结果，作为其他条件的比较基准。"],
        ["Compact Cards", "只开启紧凑卡片。所有节点仍然显示，但卡片变小，只保留类型和短标题。", "测量在不隐藏节点的情况下，卡片总面积减少多少。"],
        ["Optimized Overview", "同时开启紧凑卡片和低细节语义缩放，并将画布缩放到 50%。", "保留全部节点，但减少每张卡片的信息，用于查看整棵树的结构。"],
        ["Optimized Search", "在 Overview 状态下搜索一个预先指定且名称唯一的节点。", "突出目标、淡化其他节点，并检查目标是否进入当前视口。"],
        ["Subtree Focus", "选择一个预先指定的分支作为聚焦目标。", "只显示目标分支和必要上下文，减少当前画面中的无关节点。"],
        ["Context Collapse", "保留通往固定目标的路径，并折叠路径以外的部分子树。", "保留折叠根，用摘要代替其后代，减少同时显示的节点。"],
    ],
    "4.1": [
        ["节点", "行为"],
        ["Root", "执行唯一普通子节点。"],
        ["Sequence", "按序执行，遇到 failure 或 running 时停止，并记忆运行索引。"],
        ["Selector", "按序执行，遇到 success 或 running 时停止，并记忆运行索引。"],
        ["Reactive Selector", "每个 Tick 从高优先级重新检查，可以中断低优先级分支。"],
        ["Random Selector", "使用可设种子的随机选择，运行期间保持同一子节点。"],
        ["Parallel", "同时更新多个子节点，并应用 success 和 failure 阈值。"],
        ["Repeat", "按固定次数或无限重复，跨 Tick 保存计数。"],
        ["Wait", "累积经过时间，到达设定时长后返回 success。"],
        ["Condition", "调用 Actor 判断，或比较黑板值。"],
        ["Action", "调用 Actor 方法，并把 Bool 或状态转换为行为树状态。"],
    ],
    "5.1": [
        ["套件", "通过", "总数", "主要覆盖"],
        ["资源与运行时", "153", "153", "结构、节点语义、Decorator、Actor、Schema、调试桥和资源修改后的执行复位"],
        ["编辑器 GUI", "250", "250", "简化菜单、右键创建、图编辑、历史、Schema 与显示条件复位"],
        ["基础游戏", "13", "13", "三个敌人、移动、伤害、巡逻、死亡重启与生命周期"],
        ["复杂竞技场", "34", "34", "索敌、响应式防御、近/远程攻击、跳跃、攀爬、搜索、撤退和恢复"],
        ["合计", "450", "450", "功能与游戏断言全部通过"],
    ],
    "5.2": [
        ["指标", "31", "61", "121", "241", "364"],
        ["画布卡片", "26", "51", "101", "202", "305"],
        ["Compact 面积降低", "61.35%", "61.30%", "61.30%", "61.16%", "61.09%"],
        ["Overview 面积降低", "90.34%", "90.33%", "90.33%", "90.29%", "90.27%"],
        ["Search 淡化", "96.15%", "98.04%", "99.01%", "99.50%", "99.67%"],
        ["Focus 卡片降低", "76.92%", "86.27%", "87.13%", "91.58%", "85.90%"],
        ["Collapse 卡片降低", "69.23%", "78.43%", "89.11%", "93.07%", "96.39%"],
    ],
    "5.3": [
        ["屏幕", "表面（cm）", "画布", "241 覆盖率", "364 覆盖率", "241 图/屏幕", "364 图/屏幕"],
        ["BOE0CD1，15.94 英寸", "34×22", "1190×770", "18.32%", "12.13%", "116.38x", "184.13x"],
        ["H27T22S，26.96 英寸", "60×33", "2100×1155", "38.61%", "25.57%", "43.96x", "69.56x"],
        ["U32G3X，31.55 英寸", "70×39", "2450×1365", "46.53%", "30.82%", "31.89x", "50.45x"],
    ],
    "A.1": [
        ["条件", "开关与输入", "实际显示效果"],
        ["Baseline", "关闭实验优化，展开全部节点，清空搜索和聚焦，缩放 100%。", "完整卡片和全部分支。"],
        ["Compact Cards", "只开启 Compact。", "全部节点保留，卡片只显示类型和短标题。"],
        ["Optimized Overview", "开启 Compact 与 Semantic Zoom，缩放 50%。", "全部节点保留，使用最低细节查看整体结构。"],
        ["Optimized Search", "在 Overview 上搜索固定唯一标题。", "目标进入视口并保持显著，其他卡片淡化。"],
        ["Subtree Focus", "开启 Compact，设置固定 Focus Root。", "只显示指定分支及必要上下文。"],
        ["Context Collapse", "开启 Compact，折叠固定目标路径之外的指定分支。", "保留折叠根与摘要，隐藏其后代卡片。"],
    ],
}


TABLES_EN = {
    "3.1": [
        ["Condition", "Operation performed in the experiment", "Interface change and main measurement purpose"],
        ["Baseline", "Disable the display optimisations involved in the experiment and show full cards and every branch.", "Provide the original display as the comparison baseline for the other conditions."],
        ["Compact Cards", "Enable only compact cards. Every node remains visible, but cards become smaller and retain only the type and short title.", "Measure how much total card area is reduced without hiding nodes."],
        ["Optimized Overview", "Enable compact cards and low-detail semantic zoom together, and set canvas zoom to 50%.", "Retain every node but reduce the information on each card in order to view the structure of the complete tree."],
        ["Optimized Search", "In the Overview state, search for a predefined node with a unique name.", "Highlight the target, dim other nodes, and check whether the target enters the current viewport."],
        ["Subtree Focus", "Select a predefined branch as the focus target.", "Display only the target branch and necessary context, reducing irrelevant nodes in the current view."],
        ["Context Collapse", "Retain the path to a fixed target and collapse selected subtrees outside that path.", "Retain each collapsed root, replace its descendants with a summary, and reduce the number of simultaneously displayed nodes."],
    ],
    "4.1": [
        ["Node", "Behaviour"],
        ["Root", "Executes its single ordinary child."],
        ["Sequence", "Executes children in order, stops on failure or running, and remembers the running index."],
        ["Selector", "Executes children in order, stops on success or running, and remembers the running index."],
        ["Reactive Selector", "Rechecks from the highest priority on every Tick and can interrupt a lower-priority branch."],
        ["Random Selector", "Uses a seedable random choice and retains the same child while running."],
        ["Parallel", "Ticks multiple children and applies success and failure thresholds."],
        ["Repeat", "Repeats a fixed number of times or indefinitely and preserves the count across Ticks."],
        ["Wait", "Accumulates elapsed time and succeeds when the duration is reached."],
        ["Condition", "Calls an Actor predicate or compares a blackboard value."],
        ["Action", "Calls an Actor method and converts a Bool or status into a tree state."],
    ],
    "5.1": [
        ["Suite", "Passed", "Total", "Principal coverage"],
        ["Resource and runtime", "153", "153", "Structure, node semantics, Decorators, Actors, Schema, debug bridge and execution reset after resource modification"],
        ["Editor GUI", "250", "250", "Simplified menus, context creation, graph editing, history, Schema and display-condition reset"],
        ["Basic game", "13", "13", "Three enemies, movement, damage, patrol, death/restart and lifecycle"],
        ["Complex arena", "34", "34", "Detection, reactive defence, melee/ranged attack, jumping, climbing, search, retreat and recovery"],
        ["Total", "450", "450", "All functionality and game assertions passed"],
    ],
    "5.2": [
        ["Measure", "31", "61", "121", "241", "364"],
        ["Canvas cards", "26", "51", "101", "202", "305"],
        ["Compact area reduction", "61.35%", "61.30%", "61.30%", "61.16%", "61.09%"],
        ["Overview area reduction", "90.34%", "90.33%", "90.33%", "90.29%", "90.27%"],
        ["Search dimming", "96.15%", "98.04%", "99.01%", "99.50%", "99.67%"],
        ["Focus card reduction", "76.92%", "86.27%", "87.13%", "91.58%", "85.90%"],
        ["Collapse card reduction", "69.23%", "78.43%", "89.11%", "93.07%", "96.39%"],
    ],
    "5.3": [
        ["Screen", "Surface (cm)", "Canvas", "241 coverage", "364 coverage", "241 graph/screen", "364 graph/screen"],
        ["BOE0CD1, 15.94 inches", "34×22", "1190×770", "18.32%", "12.13%", "116.38x", "184.13x"],
        ["H27T22S, 26.96 inches", "60×33", "2100×1155", "38.61%", "25.57%", "43.96x", "69.56x"],
        ["U32G3X, 31.55 inches", "70×39", "2450×1365", "46.53%", "30.82%", "31.89x", "50.45x"],
    ],
    "A.1": [
        ["Condition", "Switches and input", "Actual display effect"],
        ["Baseline", "Disable experimental optimisations, expand every node, clear search and focus, and set zoom to 100%.", "Full cards and every branch."],
        ["Compact Cards", "Enable only Compact.", "Every node is retained; cards display only the type and short title."],
        ["Optimized Overview", "Enable Compact and Semantic Zoom and set zoom to 50%.", "Every node is retained; the overall structure is viewed at the lowest detail level."],
        ["Optimized Search", "Search for a fixed unique title in Overview.", "The target enters the viewport and remains prominent; other cards are dimmed."],
        ["Subtree Focus", "Enable Compact and set a fixed Focus Root.", "Display only the specified branch and necessary context."],
        ["Context Collapse", "Enable Compact and collapse specified branches outside the fixed target path.", "Retain each collapsed root and summary while hiding its descendant cards."],
    ],
}


CAPTIONS_ZH = {
    "t3.1": "表 3.1　六种显示实验条件及实际操作",
    "t4.1": "表 4.1　已实现的运行时节点语义",
    "t5.1": "表 5.1　最终核心自动回归",
    "t5.2": "表 5.2　五种资源规模的受控显示测量（降低比例相对 Baseline）",
    "t5.3": "表 5.3　三种物理屏幕尺寸下的复杂树覆盖率",
    "tA.1": "表 A.1　论文使用的六种显示实验条件",
    "f4.1": "图 4.1　Godot 编辑器、行为树资源、Runner 与 Actor 之间的插件架构和数据流",
    "f4.2": "图 4.2　仅位于编辑器中的活动路径和失败原因显示",
    "f4.3": "图 4.3　同一行为树的完整细节 Baseline 和 Compact Cards",
    "f4.4": "图 4.4　局部放大和结构裁剪的两种实现",
    "f4.5": "图 4.5　复杂树中的小地图、路径摘要、导航控件和自动布局",
    "f4.6": "图 4.6　连线组织和运行状态表示",
    "f5.1": "图 5.1　五种受控规模下的密度、显著性和上下文降低",
}


CAPTIONS_EN = {
    "t3.1": "Table 3.1  The six display conditions and the operations performed in the experiment",
    "t4.1": "Table 4.1  Implemented runtime node semantics",
    "t5.1": "Table 5.1  Final core automated regression",
    "t5.2": "Table 5.2  Controlled display measurements for five resource scales (reductions are relative to Baseline)",
    "t5.3": "Table 5.3  Complex-tree coverage on three physical screen sizes",
    "tA.1": "Table A.1  The six display conditions used in the dissertation",
    "f4.1": "Figure 4.1  Plugin architecture and data flow between the Godot editor, behaviour-tree resource, Runner and Actor",
    "f4.2": "Figure 4.2  Editor-only display of the active path and failure reason",
    "f4.3": "Figure 4.3  Full-detail Baseline and Compact Cards for the same behaviour tree",
    "f4.4": "Figure 4.4  Two implementations of local magnification and structural reduction",
    "f4.5": "Figure 4.5  Minimap, path summary, navigation controls and automatic layout in a complex tree",
    "f4.6": "Figure 4.6  Connection organisation and runtime-state representation",
    "f5.1": "Figure 5.1  Reductions in density, salience and context across the five controlled scales",
}


def resolve(path: Path) -> Path:
    return path if path.is_absolute() else ROOT / path


def figure_path(image_name: str) -> Path:
    try:
        path = FIGURE_PATHS[image_name]
    except KeyError as error:
        raise KeyError(f"No figure source is configured for {image_name!r}") from error
    if not path.is_file():
        raise FileNotFoundError(path)
    return path


def transform_chinese_source_text(text: str) -> str:
    if text == CHINESE_PATH_BEFORE:
        text = CHINESE_PATH_AFTER
    for fragment, addition in CROSS_REFERENCES_ZH:
        if fragment in text and addition not in text:
            text += addition
    return text


def adjust_body_text(document: Document, *, chinese: bool) -> None:
    path_before = CHINESE_PATH_BEFORE if chinese else ENGLISH_PATH_BEFORE
    path_after = CHINESE_PATH_AFTER if chinese else ENGLISH_PATH_AFTER
    cross_references = CROSS_REFERENCES_ZH if chinese else CROSS_REFERENCES_EN

    path_matches = [paragraph for paragraph in document.paragraphs if paragraph.text == path_before]
    if len(path_matches) != 1:
        raise AssertionError(f"Expected one physical-screen evidence path, found {len(path_matches)}")
    path_matches[0].text = path_after

    for fragment, addition in cross_references:
        paragraph = find_paragraph(document, fragment)
        if addition not in paragraph.text:
            paragraph.add_run(addition)


def set_east_asian(run_or_style, name: str) -> None:
    run_or_style.font.name = name
    run_or_style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def ensure_styles(document: Document, *, chinese: bool) -> None:
    for name in ("Figure Caption", "Table Caption"):
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Caption"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(9.5)
        if chinese:
            style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(6)

    try:
        front = document.styles["Front Matter Heading"]
    except KeyError:
        front = document.styles.add_style("Front Matter Heading", WD_STYLE_TYPE.PARAGRAPH)
    front.base_style = document.styles["Normal"]
    front.font.name = "Times New Roman"
    front.font.size = Pt(16)
    front.font.bold = True
    if chinese:
        front._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.first_line_indent = Pt(0)
    front.paragraph_format.space_after = Pt(12)
    ppr = front._element.get_or_add_pPr()
    for child in list(ppr):
        if child.tag == qn("w:outlineLvl"):
            ppr.remove(child)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "9")
    ppr.append(outline)


def add_complex_field(paragraph, instruction: str, placeholder: str) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(0)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, result, end):
        run = OxmlElement("w:r")
        run.append(element)
        paragraph._p.append(run)


def make_front_matter_elements(document: Document, *, chinese: bool):
    labels = (
        ("目录", "请在 Word 中更新目录。"),
        ("图目录", "请在 Word 中更新图目录。"),
        ("表目录", "请在 Word 中更新表目录。"),
    ) if chinese else (
        ("Contents", "Update this table of contents in Word."),
        ("List of Figures", "Update this list of figures in Word."),
        ("List of Tables", "Update this list of tables in Word."),
    )
    instructions = (
        'TOC \\o "1-3" \\h \\z \\u',
        'TOC \\h \\z \\t "Figure Caption,1"',
        'TOC \\h \\z \\t "Table Caption,1"',
    )
    elements = []
    for index, ((label, placeholder), instruction) in enumerate(zip(labels, instructions, strict=True)):
        heading = document.add_paragraph(label, style="Front Matter Heading")
        heading.paragraph_format.page_break_before = index > 0
        field = document.add_paragraph()
        add_complex_field(field, instruction, placeholder)
        elements.extend((heading._p, field._p))
    return elements


def insert_before(anchor, elements) -> None:
    for element in elements:
        anchor._p.addprevious(element)


def insert_after(anchor, elements) -> None:
    cursor = anchor._p
    for element in elements:
        cursor.addnext(element)
        cursor = element


def find_paragraph(document: Document, fragment: str, *, exact: bool = False):
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if (paragraph.text == fragment if exact else fragment in paragraph.text)
    ]
    if len(matches) != 1:
        raise AssertionError(f"Expected one paragraph for {fragment!r}, found {len(matches)}")
    return matches[0]


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def set_table_borders_none(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def make_table(document: Document, rows: list[list[str]], *, font_size: float) -> object:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    repeat_header(table.rows[0])
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = value
            if row_index == 0:
                shade_cell(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)
                    run.font.bold = row_index == 0
    return table


def table_block(document: Document, caption: str, rows: list[list[str]], *, font_size: float):
    caption_paragraph = document.add_paragraph(caption, style="Table Caption")
    caption_paragraph.paragraph_format.keep_with_next = True
    table = make_table(document, rows, font_size=font_size)
    return [caption_paragraph._p, table._tbl]


def single_figure_block(document: Document, image_name: str, caption: str, *, width_cm: float = 12.2):
    image_path = figure_path(image_name)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    caption_paragraph = document.add_paragraph(caption, style="Figure Caption")
    return [paragraph._p, caption_paragraph._p]


def pair_figure_block(
    document: Document,
    left_image: str,
    left_label: str,
    right_image: str,
    right_label: str,
    caption: str,
):
    elements = []
    for image_name, label in zip(
        (left_image, right_image),
        (left_label, right_label),
        strict=True,
    ):
        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.first_line_indent = Pt(0)
        image_paragraph.paragraph_format.keep_with_next = True
        image_paragraph.add_run().add_picture(str(figure_path(image_name)), width=Cm(12.0))
        label_paragraph = document.add_paragraph(label)
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_paragraph.paragraph_format.first_line_indent = Pt(0)
        label_paragraph.paragraph_format.space_after = Pt(3)
        label_paragraph.paragraph_format.keep_with_next = True
        for run in label_paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
        elements.extend((image_paragraph._p, label_paragraph._p))
    caption_paragraph = document.add_paragraph(caption, style="Figure Caption")
    elements.append(caption_paragraph._p)
    return elements


def add_all_visuals(document: Document, *, chinese: bool) -> None:
    tables = TABLES_ZH if chinese else TABLES_EN
    captions = CAPTIONS_ZH if chinese else CAPTIONS_EN
    anchors = {
        "t3.1": "本论文只比较表3.1" if chinese else "This dissertation compares only the six display conditions listed in Table 3.1.",
        "f4.1": "系统分为编辑器、资源、运行时和测试四个部分" if chinese else "The system is divided into editor, resource, runtime and test components",
        "t4.1": "表4.1概括各类节点的行为" if chinese else "Table 4.1 summarises the behaviour of each node type",
        "f4.2": "游戏画面中不会出现这些调试内容" if chinese else "None of this debugging content appears in the game view",
        "f4.3": "减少只依赖单一颜色通道的问题" if chinese else "reducing reliance on a single colour channel",
        "f4.4": "展开后仍使用原有节点位置和父子关系" if chinese else "the original node positions and parent–child relationships are still used",
        "f4.5": "自动排列仍保持同级节点从左到右的执行顺序" if chinese else "Automatic arrangement still preserves the left-to-right execution order of siblings",
        "f4.6": "实际经过的分支" if chinese else "branch actually traversed by the current Tick",
        "t5.1": "如表5.1所示" if chinese else "as shown in Table 5.1",
        "t5.2": "表5.2给出" if chinese else "Table 5.2 presents",
        "f5.1": "所有条件都保持了节点位置和从左到右的执行顺序" if chinese else "Every condition preserved node positions and left-to-right execution order",
        "t5.3": "表5.3列出" if chinese else "Table 5.3 lists",
        "tA.1": "附录 A　六种实验条件参考" if chinese else "Appendix A　Reference for the Six Experimental Conditions",
    }

    insert_after(find_paragraph(document, anchors["t3.1"]), table_block(document, captions["t3.1"], tables["3.1"], font_size=8.2))
    insert_after(find_paragraph(document, anchors["f4.1"]), single_figure_block(document, "architecture.png", captions["f4.1"], width_cm=12.5))
    insert_after(find_paragraph(document, anchors["t4.1"]), table_block(document, captions["t4.1"], tables["4.1"], font_size=8.5))
    insert_after(find_paragraph(document, anchors["f4.2"]), single_figure_block(document, "07_runtime_failure.png", captions["f4.2"], width_cm=12.0))
    insert_after(find_paragraph(document, anchors["f4.3"]), pair_figure_block(
        document, "01_baseline.png", "（a）Baseline" if chinese else "(a) Baseline",
        "02_compact.png", "（b）Compact Cards" if chinese else "(b) Compact Cards", captions["f4.3"]
    ))
    insert_after(find_paragraph(document, anchors["f4.4"]), pair_figure_block(
        document, "03_fisheye.png", "（a）鱼眼焦点与上下文" if chinese else "(a) Fisheye focus and context",
        "05_collapsed.png", "（b）子树折叠与摘要" if chinese else "(b) Subtree collapse and summary", captions["f4.4"]
    ))
    insert_after(find_paragraph(document, anchors["f4.5"]), single_figure_block(document, "10_complex_tree.png", captions["f4.5"], width_cm=12.0))
    insert_after(find_paragraph(document, anchors["f4.6"]), pair_figure_block(
        document, "08_orthogonal_edges.png", "（a）正交连线与活动路径" if chinese else "(a) Orthogonal connections and active path",
        "09_bundled_edges.png", "（b）边捆绑与失败标注" if chinese else "(b) Edge bundling and failure annotation", captions["f4.6"]
    ))
    insert_after(find_paragraph(document, anchors["t5.1"]), table_block(document, captions["t5.1"], tables["5.1"], font_size=8.0))
    insert_after(find_paragraph(document, anchors["t5.2"]), table_block(document, captions["t5.2"], tables["5.2"], font_size=7.8))
    insert_after(find_paragraph(document, anchors["f5.1"]), single_figure_block(document, "optimization_ratios.png", captions["f5.1"], width_cm=12.4))
    insert_after(find_paragraph(document, anchors["t5.3"]), table_block(document, captions["t5.3"], tables["5.3"], font_size=7.0))
    insert_after(find_paragraph(document, anchors["tA.1"], exact=True), table_block(document, captions["tA.1"], tables["A.1"], font_size=8.0))


def configure_english_document(document: Document) -> None:
    text_builder.configure_document(document)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def english_source_blocks():
    bibliography = (ENGLISH_SOURCE / "bibliography.tex").read_text(encoding="utf-8")
    citations, _ = text_builder.parse_bibliography(bibliography)
    citations["unrealbt"] = "Epic Games, n.d."
    citations["godoteditorplugin"] = "Godot Engine contributors, n.d.-a"
    citations["godotgraphedit"] = "Godot Engine contributors, n.d.-b"

    chapter_paths = [
        ENGLISH_SOURCE / "chapters" / f"{number:02d}_{name}.tex"
        for number, name in [
            (1, "introduction"),
            (2, "literature_review"),
            (3, "methodology"),
            (4, "method"),
            (5, "results"),
            (6, "discussion"),
            (7, "conclusion"),
        ]
    ]
    chapter_sources = [path.read_text(encoding="utf-8") for path in chapter_paths]
    appendix_source = (ENGLISH_SOURCE / "chapters" / "appendices.tex").read_text(encoding="utf-8")
    labels = text_builder.collect_all_labels(chapter_sources, appendix_source)

    blocks = text_builder.parse_structured_source(
        (ENGLISH_SOURCE / "frontmatter.tex").read_text(encoding="utf-8"),
        [],
        citations,
        labels,
        frontmatter=True,
    )
    for chapter_number, source in enumerate(chapter_sources, start=1):
        chapter_blocks = text_builder.parse_structured_source(
            source,
            [str(chapter_number)],
            citations,
            labels,
        )
        filtered = []
        for block in chapter_blocks:
            text = block.text
            if text.startswith("The research is divided into five stages."):
                continue
            if chapter_number == 5 and text.startswith("5.4　"):
                break
            if chapter_number == 7 and text.startswith("7.2　"):
                break
            if text.startswith("Functionality testing and the playable game confirmed that the plugin could serve"):
                continue
            text = text.replace(
                " The results can explain only geometric display differences produced by different screen sizes and cannot directly demonstrate that a large screen is easier to use.",
                "",
            )
            text = text.replace(
                " Because no human experiment was conducted, this dissertation cannot demonstrate that user understanding improved.",
                "",
            )
            text = re.sub(r"^第(\d+)章　", r"Chapter \1　", text)
            filtered.append(text_builder.Block(block.style, text))
        blocks.extend(filtered)

    appendix_blocks = text_builder.parse_structured_source(
        appendix_source,
        ["A", "B", "C"],
        citations,
        labels,
    )
    for block in appendix_blocks:
        text = re.sub(r"^附录 ([A-Z])　", r"Appendix \1　", block.text)
        if text.startswith("Appendix C　"):
            break
        blocks.append(text_builder.Block(block.style, text))

    architecture_fragment = "The system is divided into editor, resource, runtime and test components"
    architecture_matches = [index for index, block in enumerate(blocks) if architecture_fragment in block.text]
    if len(architecture_matches) != 1:
        raise AssertionError(f"Expected one architecture paragraph, found {len(architecture_matches)}")
    index = architecture_matches[0]
    blocks[index] = text_builder.Block(
        blocks[index].style,
        blocks[index].text
        + " The editor interface is integrated through Godot's EditorPlugin extension mechanism, "
        + "and its node canvas is based on GraphEdit (Godot Engine contributors, n.d.-a, n.d.-b).",
    )
    return blocks


def write_reference_paragraph(document: Document, prefix: str, url: str, suffix: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(prefix)
    add_hyperlink(paragraph, url, url)
    paragraph.add_run(suffix)


def build_chinese(input_path: Path, output_path: Path) -> None:
    shutil.copy2(input_path, output_path)
    document = Document(output_path)
    ensure_styles(document, chinese=True)
    original_texts = [paragraph.text for paragraph in document.paragraphs]
    acknowledgement = find_paragraph(document, "致谢", exact=True)
    insert_before(acknowledgement, make_front_matter_elements(document, chinese=True))
    adjust_body_text(document, chinese=True)
    add_all_visuals(document, chinese=True)
    document.core_properties.subject = "中文论文完整 DOCX（图表与引用齐全）"
    document.core_properties.comments = (
        "Built from the accepted citation-verified text DOCX; figures and tables restored without PDF extraction."
    )
    document.save(output_path)
    verify_complete_docx(output_path, chinese=True, expected_headings=49)

    actual = [paragraph.text for paragraph in Document(output_path).paragraphs]
    cursor = 0
    for original in original_texts:
        expected = transform_chinese_source_text(original)
        while cursor < len(actual) and actual[cursor] != expected:
            cursor += 1
        if cursor >= len(actual):
            raise AssertionError(f"Chinese source paragraph was not preserved: {expected[:100]!r}")
        cursor += 1


def build_english(output_path: Path) -> None:
    document = Document()
    configure_english_document(document)
    ensure_styles(document, chinese=False)
    document.core_properties.title = (
        "Display Optimisation and Experimental Evaluation of Large Visual Behaviour Trees in Godot 4.6"
    )
    document.core_properties.subject = "Supervisor review draft without acknowledgements or declaration"
    document.core_properties.author = "[AUTHOR NAME]"
    document.core_properties.comments = (
        "English supervisor-review DOCX derived from the accepted Chinese scope and tracked English translation; no PDF extraction."
    )

    title = document.add_paragraph(document.core_properties.title, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(30)
    for value in [
        "by",
        "[AUTHOR NAME]",
        "A thesis submitted in partial fulfilment",
        "of the requirements for the degree of",
        "Master of Science in Games Engineering",
        "University of Warwick",
        "WMG",
        "September 2026",
    ]:
        paragraph = document.add_paragraph(value)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
    document.add_page_break()
    make_front_matter_elements(document, chinese=False)

    blocks = english_source_blocks()
    reference_inserted = False
    for block in blocks:
        if block.text.startswith("Appendix A　") and not reference_inserted:
            heading = document.add_paragraph("References", style="Heading 1")
            heading.paragraph_format.page_break_before = True
            for prefix, url, suffix in REFERENCES:
                write_reference_paragraph(document, prefix, url, suffix)
            reference_inserted = True
        style = "Normal" if block.style == "Formula" else block.style
        paragraph = document.add_paragraph(block.text, style=style)
        if block.style.startswith("Heading"):
            paragraph.paragraph_format.page_break_before = block.style == "Heading 1"
        elif block.style == "Formula":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
    if not reference_inserted:
        raise AssertionError("English references were not inserted")

    adjust_body_text(document, chinese=False)
    add_all_visuals(document, chinese=False)
    document.save(output_path)
    verify_complete_docx(output_path, chinese=False, expected_headings=46)


def external_hyperlinks(document: Document) -> list[str]:
    return [
        relationship.target_ref
        for relationship in document.part.rels.values()
        if relationship.reltype.endswith("/hyperlink") and relationship.is_external
    ]


def verify_complete_docx(path: Path, *, chinese: bool, expected_headings: int) -> None:
    document = Document(path)
    if len(document.tables) != 6:
        raise AssertionError(f"Expected 6 data tables, found {len(document.tables)}")
    if len(document.inline_shapes) != 10:
        raise AssertionError(f"Expected 10 images, found {len(document.inline_shapes)}")
    figure_captions = [p.text for p in document.paragraphs if p.style.name == "Figure Caption"]
    table_captions = [p.text for p in document.paragraphs if p.style.name == "Table Caption"]
    if len(figure_captions) != 7 or len(table_captions) != 6:
        raise AssertionError(
            f"Caption mismatch: figures={len(figure_captions)}, tables={len(table_captions)}"
        )
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    if len(headings) != expected_headings:
        raise AssertionError(f"Expected {expected_headings} headings, found {len(headings)}")

    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
    front_labels = ("目录", "图目录", "表目录") if chinese else ("Contents", "List of Figures", "List of Tables")
    first_body_label = "致谢" if chinese else "Abstract"
    front_indices = []
    for label in front_labels:
        matches = [index for index, text in enumerate(paragraph_texts) if text == label]
        if len(matches) != 1:
            raise AssertionError(f"Expected one front-matter heading {label!r}, found {len(matches)}")
        front_indices.append(matches[0])
    body_matches = [index for index, text in enumerate(paragraph_texts) if text == first_body_label]
    if len(body_matches) != 1:
        raise AssertionError(f"Expected one first body heading {first_body_label!r}, found {len(body_matches)}")
    if front_indices != sorted(front_indices) or front_indices[-1] >= body_matches[0]:
        raise AssertionError(
            f"Front-matter order is invalid: front={front_indices}, first_body={body_matches[0]}"
        )

    xml_text = path.read_bytes()
    with ZipFile(path) as archive:
        if archive.testzip() is not None:
            raise AssertionError("DOCX ZIP validation failed")
        document_xml = archive.read("word/document.xml").decode("utf-8")
    if document_xml.count("TOC ") != 3:
        raise AssertionError("Expected three TOC/list fields")

    text = "\n".join(p.text for p in document.paragraphs)
    if chinese:
        required = ("表 3.1", "图 4.1", "图 5.1", "表 A.1", "https://doi.org/10.1145/22627.22342")
    else:
        required = (
            "Chapter 1　Introduction",
            "Appendix A　Reference for the Six Experimental Conditions",
            "Figure 4.1",
            "Table 5.3",
            "(Epic Games, n.d.)",
            "(Godot Engine contributors, n.d.-a, n.d.-b)",
        )
        forbidden = (
            "Acknowledgements",
            "Declaration",
            "Draft Completion Checklist",
            "5.4　Results for the Research Question",
            "7.2　Contributions",
            "7.3　Limitations",
            "7.4　Future Work",
            "7.5　Final Conclusion",
        )
        for fragment in forbidden:
            if fragment in text:
                raise AssertionError(f"Supervisor-only English DOCX retained forbidden text: {fragment}")
        cjk = re.search(r"[\u3400-\u9fff]", text)
        if cjk:
            raise AssertionError(f"English DOCX contains Chinese text near: {text[max(0, cjk.start()-30):cjk.start()+30]!r}")
        links = external_hyperlinks(document)
        if len(links) != len(URLS) or set(links) != set(URLS):
            raise AssertionError(f"English bibliography hyperlink mismatch: {links}")
    for fragment in required:
        if fragment not in text:
            raise AssertionError(f"Required text missing: {fragment}")
    if not xml_text:
        raise AssertionError("DOCX is empty")


def print_summary(label: str, path: Path) -> None:
    document = Document(path)
    digest = hashlib.sha256(path.read_bytes()).hexdigest().upper()
    print(f"{label}_OUTPUT={path}")
    print(f"{label}_SHA256={digest}")
    print(f"{label}_PARAGRAPHS={len(document.paragraphs)}")
    print(f"{label}_HEADINGS={sum(p.style.name.startswith('Heading') for p in document.paragraphs)}")
    print(f"{label}_TABLES={len(document.tables)}")
    print(f"{label}_IMAGES={len(document.inline_shapes)}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--chinese-input", required=True, type=Path)
    parser.add_argument("--chinese-output", required=True, type=Path)
    parser.add_argument("--english-output", required=True, type=Path)
    args = parser.parse_args()

    chinese_input = resolve(args.chinese_input).resolve()
    chinese_output = resolve(args.chinese_output).resolve()
    english_output = resolve(args.english_output).resolve()
    if chinese_input in (chinese_output, english_output) or chinese_output == english_output:
        raise ValueError("Input and output paths must be distinct")
    if not chinese_input.is_file():
        raise FileNotFoundError(chinese_input)
    chinese_output.parent.mkdir(parents=True, exist_ok=True)
    english_output.parent.mkdir(parents=True, exist_ok=True)

    build_chinese(chinese_input, chinese_output)
    build_english(english_output)
    print_summary("CHINESE", chinese_output)
    print_summary("ENGLISH", english_output)
    print("FULL_DOCX_BUILD=PASS")


if __name__ == "__main__":
    main()
