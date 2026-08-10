from __future__ import annotations

import csv
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "testgame" / "testgame"
RESULTS = PROJECT / "test_results"
VISUAL = RESULTS / "visual"
HUMAN_VISUAL = RESULTS / "human_study_visual"
ASSETS = ROOT / "docs" / "report_assets"
OUTPUT = ROOT / "docs" / "Godot_Behavior_Tree_Plugin_Complete_Report.docx"

FONT_NAME = "Microsoft YaHei"
FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")
NAVY = "17324D"
BLUE = "2E74B5"
TEAL = "159A9C"
GOLD = "D99B2B"
INK = "1E2933"
MUTED = "637282"
LIGHT = "EEF3F7"
LIGHT_TEAL = "E7F6F5"
LIGHT_GOLD = "FFF5DC"
WHITE = "FFFFFF"
RED = "C94A4A"
GREEN = "31865B"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, size: float | None = None, color: str = INK,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120,
                     bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths: list[float], font_size: float = 8.5,
              align_columns: set[int] | None = None) -> object:
    align_columns = align_columns or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, font_size, WHITE, True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            if row_index % 2 == 1:
                set_cell_shading(cell, "F7F9FB")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in align_columns else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_run_font(run, font_size, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.17


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])


def add_running_furniture(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("GODOT 4.6  ·  VISUAL BEHAVIOR TREE")
        set_run_font(run, 8.5, MUTED, True)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_after = Pt(0)
        run = fp.add_run("Godot 4.6 可视化行为树插件完整报告  |  第 ")
        set_run_font(run, 8, MUTED)
        add_page_number(fp)
        run = fp.add_run(" 页")
        set_run_font(run, 8, MUTED)


def add_title(doc: Document, text: str, size: float = 24, color: str = NAVY,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, after: float = 8) -> None:
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size, color, True)


def add_kicker(doc: Document, text: str, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_font(run, 9.5, TEAL, True)


def add_body(doc: Document, text: str, bold_prefix: str | None = None,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = alignment
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, 10.5, INK, True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, 10.5, INK)
    else:
        run = p.add_run(text)
        set_run_font(run, 10.5, INK)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 10.25, INK)


def add_number(doc: Document, text: str) -> None:
    numbering = doc.part.numbering_part.element
    last_is_numbered = bool(doc.paragraphs and doc.paragraphs[-1]._p.pPr is not None
                            and doc.paragraphs[-1]._p.pPr.numPr is not None)
    if not last_is_numbered:
        abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
        num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
        abstract_id = max(abstract_ids, default=0) + 1
        num_id = max(num_ids, default=0) + 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), "%1.")
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "720")
        indent.set(qn("w:hanging"), "360")
        p_pr.extend([tabs, indent])
        level.extend([start, num_fmt, level_text, suffix, p_pr])
        abstract.append(level)
        first_num_index = next((index for index, child in enumerate(numbering)
                                if child.tag == qn("w:num")), len(numbering))
        numbering.insert(first_num_index, abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        level_override = OxmlElement("w:lvlOverride")
        level_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        level_override.append(start_override)
        num.append(level_override)
        numbering.append(num)
        doc._bt_number_id = num_id
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(doc._bt_number_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.17
    run = p.add_run(text)
    set_run_font(run, 10.25, INK)


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_TEAL,
                accent: str = TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(label + "  ")
    set_run_font(run, 10.5, accent, True)
    run = p.add_run(text)
    set_run_font(run, 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = False
    run = p.add_run(text)
    set_run_font(run, 8.5, MUTED, italic=True)


def add_picture(doc: Document, path: Path, width: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_two_pictures(doc: Document, left: Path, right: Path,
                     left_caption: str, right_caption: str) -> None:
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, [3.2, 3.2])
    for index, path in enumerate((left, right)):
        cell = table.cell(0, index)
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(3.05))
    for index, text in enumerate((left_caption, right_caption)):
        cell = table.cell(1, index)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, 8.2, MUTED, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def read_display_data() -> dict[int, dict[str, dict[str, float]]]:
    result: dict[int, dict[str, dict[str, float]]] = {}
    with (RESULTS / "display_optimization_raw.csv").open(encoding="utf-8-sig") as file:
        for row in csv.DictReader(file):
            size = int(row["tree_size"])
            result.setdefault(size, {})[row["method"]] = {
                key: float(value) for key, value in row.items()
                if key not in ("tree_size", "method")
            }
    return result


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig") as file:
        return list(csv.DictReader(file))


def create_charts(display: dict[int, dict[str, dict[str, float]]],
                  branch_rows: list[dict[str, str]], minimap_rows: list[dict[str, str]]) -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    regular = lambda size: ImageFont.truetype(str(FONT_PATH), size)
    bold = lambda size: ImageFont.truetype(str(Path(r"C:\Windows\Fonts\msyhbd.ttc")), size)

    image = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((900, 58), "插件总体架构与数据流", font=bold(44), fill=f"#{NAVY}", anchor="mm")
    boxes = [
        (70, 180, 360, 190, "Godot 编辑器插件", "节点创建、拖拽、连接\n布局与显示优化", BLUE),
        (525, 180, 360, 190, "BTTreeResource", "树结构、节点资源\n校验、保存与复制", TEAL),
        (980, 180, 360, 190, "BehaviorTreeRunner", "状态机、节点记忆\n黑板与 Decorator", GOLD),
        (1450, 180, 280, 190, "Actor / NPC", "Action 与 Condition\n方法调用", GREEN),
        (525, 540, 360, 190, "Live Debug Bridge", "原子 JSON 快照\n活动路径与失败原因", RED),
        (980, 540, 360, 190, "测试游戏", "巡逻、追击、攻击\n搜索、撤退、治疗", NAVY),
    ]
    for x, y, w, h, title, detail, color in boxes:
        draw.rounded_rectangle((x, y, x + w, y + h), radius=24, fill="#F7FAFC", outline=f"#{color}", width=4)
        draw.text((x + w / 2, y + 56), title, font=bold(28), fill=f"#{color}", anchor="mm")
        for line_index, line in enumerate(detail.splitlines()):
            draw.text((x + w / 2, y + 118 + line_index * 35), line, font=regular(21), fill=f"#{INK}", anchor="mm")
    arrows = [
        ((430, 275), (525, 275)), ((885, 275), (980, 275)),
        ((1340, 275), (1450, 275)), ((1160, 370), (1160, 540)),
        ((980, 635), (885, 635)), ((705, 540), (705, 370)),
    ]
    for start, end in arrows:
        draw.line((start, end), fill="#718096", width=5)
        ex, ey = end
        sx, sy = start
        if ex != sx:
            sign = 1 if ex > sx else -1
            draw.polygon([(ex, ey), (ex - 18 * sign, ey - 11), (ex - 18 * sign, ey + 11)], fill="#718096")
        else:
            sign = 1 if ey > sy else -1
            draw.polygon([(ex, ey), (ex - 11, ey - 18 * sign), (ex + 11, ey - 18 * sign)], fill="#718096")
    draw.text((900, 835), "编辑器负责可视化建模，Runner 负责运行时语义，Live Debug 将运行状态回传到编辑器。",
              font=regular(22), fill=f"#{MUTED}", anchor="mm")
    image.save(ASSETS / "architecture.png")

    large = display[364]
    baseline = large["Baseline"]
    labels = ["Compact\n卡片面积", "Semantic\n信息字段", "Subtree Focus\n可见节点", "Collapse\n可见节点", "Search\n降噪节点"]
    values = [
        (1 - large["Compact Cards"]["card_area_px2"] / baseline["card_area_px2"]) * 100,
        (1 - large["Semantic Zoom"]["information_fields"] / baseline["information_fields"]) * 100,
        (1 - large["Subtree Focus"]["visible_ratio"]) * 100,
        (1 - large["Subtree Collapse"]["visible_ratio"]) * 100,
        large["Search Highlight"]["dimmed_nodes"] / 364 * 100,
    ]
    image = Image.new("RGB", (1720, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((860, 55), "364 节点行为树显示优化效果", font=bold(42), fill=f"#{NAVY}", anchor="mm")
    plot_left, plot_top, plot_right, plot_bottom = 140, 130, 1650, 720
    draw.line((plot_left, plot_top, plot_left, plot_bottom), fill="#7B8794", width=3)
    draw.line((plot_left, plot_bottom, plot_right, plot_bottom), fill="#7B8794", width=3)
    for tick in range(0, 101, 20):
        y = plot_bottom - tick / 100 * (plot_bottom - plot_top)
        draw.line((plot_left, y, plot_right, y), fill="#DFE5EA", width=2)
        draw.text((plot_left - 22, y), f"{tick}", font=regular(20), fill=f"#{MUTED}", anchor="rm")
    colors = [TEAL, BLUE, GOLD, RED, NAVY]
    bar_width = 190
    gap = 95
    start_x = 235
    for index, (label, value, color) in enumerate(zip(labels, values, colors)):
        x0 = start_x + index * (bar_width + gap)
        y0 = plot_bottom - value / 100 * (plot_bottom - plot_top)
        draw.rounded_rectangle((x0, y0, x0 + bar_width, plot_bottom), radius=12, fill=f"#{color}")
        draw.text((x0 + bar_width / 2, y0 - 30), f"{value:.1f}%", font=bold(25), fill=f"#{INK}", anchor="mm")
        lines = label.split("\n")
        for li, line in enumerate(lines):
            draw.text((x0 + bar_width / 2, 770 + li * 32), line, font=regular(21), fill=f"#{INK}", anchor="mm")
    draw.text((42, 420), "优化/降噪比例（%）", font=regular(23), fill=f"#{MUTED}", anchor="mm")
    image.save(ASSETS / "optimization_ratios.png")

    sizes = [int(row["tree_size"]) for row in branch_rows]
    branch_first = [float(row["first_update_ms"]) for row in branch_rows]
    branch_steady_us = [float(row["steady_update_ms"]) * 1000 for row in branch_rows]
    minimap_toggle = [float(row["toggle_cycle_ms"]) for row in minimap_rows]
    minimap_steady_us = [float(row["steady_status_update_ms"]) * 1000 for row in minimap_rows]
    image = Image.new("RGB", (1800, 820), "white")
    draw = ImageDraw.Draw(image)
    draw.text((900, 48), "显示功能在不同树规模下的性能", font=bold(42), fill=f"#{NAVY}", anchor="mm")

    def line_chart(rect, title, ylabel, series):
        left, top, right, bottom = rect
        draw.text(((left + right) / 2, top - 38), title, font=bold(28), fill=f"#{INK}", anchor="mm")
        draw.line((left, top, left, bottom), fill="#7B8794", width=3)
        draw.line((left, bottom, right, bottom), fill="#7B8794", width=3)
        max_value = max(max(values) for _, values, _ in series) * 1.15
        for step in range(5):
            value = max_value * step / 4
            y = bottom - step / 4 * (bottom - top)
            draw.line((left, y, right, y), fill="#E2E7EC", width=2)
            draw.text((left - 15, y), f"{value:.1f}", font=regular(17), fill=f"#{MUTED}", anchor="rm")
        xs = [left + (right - left) * i / (len(sizes) - 1) for i in range(len(sizes))]
        for x, size in zip(xs, sizes):
            draw.text((x, bottom + 28), str(size), font=regular(19), fill=f"#{INK}", anchor="mm")
        for label, values, color in series:
            points = [(x, bottom - value / max_value * (bottom - top)) for x, value in zip(xs, values)]
            draw.line(points, fill=f"#{color}", width=5)
            for point in points:
                x, y = point
                draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill=f"#{color}")
        legend_x = left + 15
        for index, (label, _, color) in enumerate(series):
            y = bottom + 72 + index * 31
            draw.line((legend_x, y, legend_x + 38, y), fill=f"#{color}", width=5)
            draw.text((legend_x + 50, y), label, font=regular(18), fill=f"#{INK}", anchor="lm")
        draw.text((left, top - 18), f"单位：{ylabel}", font=regular(18), fill=f"#{MUTED}", anchor="lm")

    line_chart((120, 150, 840, 610), "一次性更新成本", "毫秒", [
        ("分支淡化首次更新", branch_first, RED), ("小地图开关周期", minimap_toggle, BLUE)])
    line_chart((1020, 150, 1740, 610), "稳态状态刷新成本", "微秒", [
        ("分支淡化稳态", branch_steady_us, TEAL), ("小地图状态稳态", minimap_steady_us, GOLD)])
    image.save(ASSETS / "performance_scaling.png")


def add_cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(66)
    add_kicker(doc, "毕业设计 · 项目验收与技术总结", WD_ALIGN_PARAGRAPH.CENTER)
    add_title(doc, "Godot 4.6 可视化行为树插件", 28, NAVY, WD_ALIGN_PARAGRAPH.CENTER, 6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("完整功能、使用方法、测试与显示优化报告")
    set_run_font(run, 15, BLUE, False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run("面向 2D NPC 可视化编程、运行时执行与编辑器 Live Debug")
    set_run_font(run, 10.5, MUTED, italic=True)
    add_picture(doc, ASSETS / "architecture.png", 6.1, "系统由编辑器、资源模型、运行器、Actor 与调试桥构成")
    rows = [
        ["开发平台", "Godot 4.6 stable / Windows"],
        ["插件名称", "Behavior Tree Editor"],
        ["插件版本", "0.9.0（可分发毕业设计候选版）"],
        ["测试规模", "447 项核心断言 + 511 项性能断言 + 53 项发布包验证"],
        ["报告日期", "2026 年 8 月 9 日"],
    ]
    add_table(doc, ["项目项", "内容"], rows, [1.45, 5.05], 9, {0})


def add_contents(doc: Document) -> None:
    add_kicker(doc, "REPORT GUIDE")
    add_title(doc, "报告导读", 23, NAVY, after=10)
    add_callout(doc, "核心结论", "插件已经形成“可视化建模—资源保存—运行时执行—编辑器调试—复杂场景验证—洁净安装”的完整闭环。447 项核心自动化断言全部通过；在 364 节点基准树上，紧凑模式减少 55.9% 卡片面积，子树聚焦减少 97.5% 可见节点，子树折叠减少 99.5% 可见节点。")
    doc.add_heading("目录", level=1)
    contents = [
        ("1", "项目概述与目标"), ("2", "总体架构与实现"), ("3", "已完成功能"),
        ("4", "插件安装与具体使用方法"), ("5", "示例游戏与验证场景"),
        ("6", "测试方法与结果"), ("7", "显示优化实验与量化数据"),
        ("8", "问题修复、局限与后续工作"), ("9", "结论"), ("附录", "资源、参数与复现命令"),
    ]
    for number, title in contents:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(f"{number}  ")
        set_run_font(run, 10.5, TEAL, True)
        run = p.add_run(title)
        set_run_font(run, 10.5, INK)
    doc.add_heading("关键指标一览", level=1)
    add_table(doc, ["指标", "结果", "说明"], [
        ["核心自动化断言", "447 / 447", "资源、运行时、编辑器、游戏集成与视觉回归"],
        ["显示功能开关", "19 / 19", "均可独立启停、保存状态并验证关闭后无可见残留"],
        ["真实视觉截图", "21 张", "17 张核心 + 4 张受控实验，1600×900 GPU 渲染"],
        ["最大自动基准树", "364 节点", "同时测试 31、121、364 三档规模"],
        ["示例行为树", "4 个", "基础巡逻战斗、演示 NPC、守卫示例、复杂守卫"],
    ], [1.45, 1.25, 3.8], 9, {1})


def add_overview(doc: Document) -> None:
    doc.add_heading("1 项目概述与目标", level=1)
    add_body(doc, "本项目针对 Godot 4.6 开发一个基于 2D GUI 的可视化行为树插件，使开发者可以在 Godot 编辑器内完成 NPC 决策逻辑的创建、连接、保存、执行和调试。项目定位不是复刻 Unreal Engine 的全部 AI 系统，而是在毕业设计范围内实现一个功能闭环完整、可以真实驱动 NPC、具备大树可读性优化并有可复现实验数据的行为树组件。")
    doc.add_heading("1.1 设计目标", level=2)
    add_bullet(doc, "可视化编程：节点以上下端口连接，树从上向下展开，Sequence 等多子节点按横向从左到右执行。")
    add_bullet(doc, "运行时可用：Action 能调用指定 Actor 方法，Condition 和 Decorator 能读取黑板并影响执行结果。")
    add_bullet(doc, "可组件化绑定：NPC 通过 BehaviorTreeComponent 获得行为树，支持父节点默认 Actor、ActorPath 和显式 set_actor。")
    add_bullet(doc, "编辑器调试：游戏运行时仅在 Godot 行为树面板显示当前路径、叶节点状态和失败原因。")
    add_bullet(doc, "大树可读性：每种显示优化均有独立开关，功能关闭后恢复基础状态，避免实验性功能相互污染。")
    add_bullet(doc, "可验证性：使用资源、运行时、GUI、集成场景、烟雾测试、性能基准和真实截图构成多层测试。")
    doc.add_heading("1.2 当前完成范围", level=2)
    add_body(doc, "已完成的范围包括 10 种节点类型、资源校验、编辑操作历史、拖拽与单线连接、附属及结构型 Decorator、类型化 Inspector、Blackboard Schema Editor、运行时状态机、编辑器 Live Debug 与 Live Blackboard、19 项显示优化、基础 2D 游戏和复杂守卫竞技场。Service 节点、子树资源复用和完整 UE 风格 AI 框架不在当前范围内。")
    add_callout(doc, "完成标准", "NPC 的行为必须由行为树资源与 BehaviorTreeRunner 实际驱动，而不是由测试脚本模拟；编辑器中的高亮必须来自运行时快照，而不是游戏内叠加 UI。", LIGHT_GOLD, GOLD)

    doc.add_heading("2 总体架构与实现", level=1)
    add_picture(doc, ASSETS / "architecture.png", 6.35, "图 1  插件总体架构及运行数据流")
    add_body(doc, "编辑器插件通过 EditorPlugin 将 BTEditorView 加入 Godot 底部面板；BTEditorView 使用 GraphEdit、GraphNode 和标准 Control 构建节点画布、工具栏、节点面板和 Inspector。BTTreeResource 与 BTNodeResource 保存拓扑、位置和参数。BehaviorTreeRunner 在游戏进程中解释资源并调用 Actor。Live Debug Bridge 使用 JSON 快照将活动路径、状态、黑板和失败原因回传编辑器。")
    doc.add_heading("2.1 核心模块", level=2)
    add_table(doc, ["模块", "职责", "关键实现"], [
        ["plugin.gd", "插件生命周期", "进入编辑器时创建底部面板，退出时安全释放"],
        ["BTEditorView", "编辑与调试界面", "节点创建、连接、Inspector、历史、显示开关、Live Debug"],
        ["BTGraphEdit / BTGraphNode", "图形画布与节点卡片", "端口、拖拽、连线、折叠预览、状态样式"],
        ["BTTreeResource", "树资源模型", "查找、排序、深复制、结构校验、Schema 与 Decorator 归属"],
        ["BehaviorTreeRunner", "运行时执行", "状态传播、节点记忆、类型黑板、Actor 调用、调试快照"],
        ["BehaviorTreeComponent", "NPC 组件入口", "继承 Runner，作为场景树中的可配置组件"],
    ], [1.45, 2.0, 3.05], 8.5)
    doc.add_heading("2.2 执行状态与顺序", level=2)
    add_body(doc, "节点统一返回 SUCCESS、FAILURE 或 RUNNING。Sequence 依次执行子节点，遇到 FAILURE 立即失败，遇到 RUNNING 保存当前索引；Selector 遇到 SUCCESS 即成功。Reactive Selector 每帧从高优先级分支重新判断，可以让攻击或追击及时打断巡逻。Parallel 同一 tick 推进多个未完成子节点，Random Selector 以可选种子生成本轮顺序，Repeat 和 Wait 保留各自运行记忆。子节点以横坐标排序，横坐标相同时使用纵坐标作为稳定次序，因此画布上的从左到右顺序就是实际执行顺序。")


def add_features(doc: Document) -> None:
    doc.add_heading("3 已完成功能", level=1)
    doc.add_heading("3.1 节点类型与运行语义", level=2)
    add_table(doc, ["节点", "子节点限制", "运行语义", "主要参数"], [
        ["Root", "最多 1 个", "行为树唯一入口；空 Root 返回 SUCCESS", "无"],
        ["Sequence", "多个", "从左到右执行，失败即停，支持 RUNNING 恢复", "无"],
        ["Selector", "多个", "从左到右选择首个成功分支", "reactive"],
        ["Random Selector", "多个", "每轮随机尝试子节点；固定种子可复现", "seed"],
        ["Parallel", "多个", "同一 tick 推进所有未完成子节点", "success_policy / failure_policy"],
        ["Repeat", "最多 1 个", "有限次数重复或无限循环；失败向上传播", "repeat_count"],
        ["Action", "无", "调用 Actor 的具体行为方法", "action_name"],
        ["Condition", "无", "调用条件方法或使用运算符比较黑板值", "method / key / operator / value"],
        ["Wait", "无", "按 delta 累积等待，不调用 Actor", "duration"],
        ["Decorator", "结构型最多 1 个；附属型不占连线", "阻断、反转、重复、限时、冷却或强制结果", "mode 及模式参数"],
    ], [0.85, 1.05, 2.85, 1.75], 8.2, {0, 1})
    doc.add_heading("3.2 Action、Condition 与 Actor", level=2)
    add_body(doc, "Action 的 parameters.action_name 指定 Actor 上的方法名；Condition 可使用 parameters.condition_name 调用 Actor 方法，也可以直接比较黑板值。方法接收 blackboard、delta 和节点资源，返回 BTStatus 或布尔值。Runner 默认把父节点作为 Actor，也支持 actor_path 或 set_actor(actor) 显式绑定。显式 Actor 在 start_tree() 后仍保持优先级。")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    set_cell = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F4F7")
    set_cell.append(shd)
    code = "func patrol(blackboard: Dictionary, delta: float, node: Resource) -> int:\n    return BTStatus.RUNNING"
    run = p.add_run(code)
    run.font.name = "Consolas"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(NAVY)

    doc.add_heading("3.3 黑板与 Decorator", level=2)
    add_body(doc, "每个 BehaviorTreeRunner 提供独立 Dictionary 黑板，可由 Actor 在运行中写入 player_detected、health、last_known_player_x 等数据。BTBlackboardSchema 资源为稳定键定义 Bool、Int、Float、String 或 Vector2 类型、默认值和说明；Runner 补齐缺失默认值并在调试快照中报告类型错误。allow_dynamic_keys 默认开启，以兼容 Action 的临时计时键。")
    add_table(doc, ["类别", "模式/运算符", "用途"], [
        ["附属 Decorator", "blackboard", "按 exists、not_exists、is_true、is_false、==、!=、>、<、>=、<= 判断黑板"],
        ["附属 Decorator", "cooldown / time_limit", "控制重试间隔和 Action 最大运行时间"],
        ["附属 Decorator", "force_success / force_failure", "在前置判断层强制结果"],
        ["结构型 Decorator", "invert / repeat_forever", "反转子节点结果或持续重复子树"],
        ["结构型 Decorator", "always_success / always_failure", "包装一个子节点并覆盖最终结果"],
        ["通用能力", "enabled / failure reason", "禁用节点安全失败，并在 Live Debug 中给出原因"],
    ], [1.25, 2.15, 3.1], 8.4)

    doc.add_heading("3.4 可视化编辑能力", level=2)
    add_bullet(doc, "节点面板拖拽创建、画布右键创建、节点右键菜单和上下方形连接端口。")
    add_bullet(doc, "连接、断开、防自连接、防循环、叶节点限制、Root 唯一性和结构校验。")
    add_bullet(doc, "节点拖拽位置持久化；添加、删除、连接、断开、拖动和 Inspector 编辑支持撤销/重做。")
    add_bullet(doc, "删除父节点时同时删除整个子树和附属 Decorator；保存前检查悬空节点和非法结构。")
    add_bullet(doc, "行为树资源选择器可直接扫描项目中的 .tres，减少手动输入精确路径的负担。")
    add_bullet(doc, "Inspector 按节点类型显示方法、模式、运算符、策略、种子、次数和持续时间控件；Advanced JSON 保留自定义参数兼容。")
    add_picture(doc, VISUAL / "01_baseline.png", 6.35, "图 2  行为树编辑器基础界面：节点面板、GraphEdit、Inspector、小地图与路径栏")


def add_display_features(doc: Document) -> None:
    doc.add_heading("3.5 十九项显示优化功能", level=2)
    rows = [
        ["Fisheye / Focus+Context", "开", "鼠标邻域内部放大，保持全局上下文；最大 1.20 倍"],
        ["Subtree Collapse / Expand", "开", "折叠子树并显示隐藏节点数量和两层摘要"],
        ["Compact Mode", "关", "缩小卡片，仅保留类型颜色和短标题"],
        ["Active Path Highlight", "开", "高亮运行路径与当前叶节点状态"],
        ["Non-active Branch Dimming", "关", "运行时淡化不在当前路径上的分支"],
        ["Multi-column Layout", "关", "宽扇出分支分列布局且保持执行顺序"],
        ["Enhanced Minimap", "开", "230×150 总览，显示覆盖节点和当前缩放"],
        ["Semantic Zoom", "关", "缩小时减少信息字段，不改变节点基础几何"],
        ["Path Summary View", "开", "显示可点击的 Actor、状态、深度和运行链条"],
        ["Decorator Condition Badges", "开", "在所有者节点上显示 Decorator 条件摘要"],
        ["Search + Highlight", "开", "按标题、类型、描述、方法和 Decorator 参数搜索"],
        ["Orthogonal Edges", "关", "使用直角折线降低曲线交叉的视觉干扰"],
        ["Edge Bundling", "关", "同方向连线共享中间主干，压缩密集连线"],
        ["Stable Incremental Layout", "关", "自动布局时尽量保留已有无冲突位置"],
        ["Breadcrumb Navigation", "开", "显示选中节点的祖先层级并可快速定位"],
        ["Failure Reason Annotation", "开", "节点红色失败标签与可点击失败摘要菜单"],
        ["Shape / Icon Type Encoding", "开", "使用形状和图标补充颜色编码，缩小与灰度下仍可辨认类型"],
        ["Accessibility / Colorblind Palette", "关", "切换色盲安全配色，并提供键盘搜索与结果导航"],
        ["Single Connection Rendering", "开", "每条关系只显示一条自定义连线，并保留命中、断开与撤销重做"],
    ]
    add_table(doc, ["功能", "默认", "作用"], rows, [2.1, 0.55, 3.85], 7.8, {1})
    add_callout(doc, "独立开关原则", "全部显示功能都保存在 user://behavior_tree_editor_view.cfg。关闭任一功能不会修改行为树执行数据；组合启用后再全部关闭的可见残留测试已经通过。")
    add_two_pictures(doc, VISUAL / "03_fisheye.png", VISUAL / "06_search.png",
                     "图 3a  鱼眼局部放大", "图 3b  搜索结果高亮与非匹配降噪")


def add_usage(doc: Document) -> None:
    doc.add_heading("4 插件安装与具体使用方法", level=1)
    doc.add_heading("4.1 安装与启用", level=2)
    add_number(doc, "将 addons/behavior_tree_editor 整个文件夹放到目标 Godot 项目的 res://addons/ 下。")
    add_number(doc, "打开 Project > Project Settings > Plugins。")
    add_number(doc, "找到 Behavior Tree Editor，将 Status 切换为 Enable。")
    add_number(doc, "回到 Godot 主界面，在底部面板点击 Behavior Tree。")
    add_callout(doc, "项目内现成路径", "当前测试项目已安装在 testgame/testgame/addons/behavior_tree_editor；visual_scripting/addons/behavior_tree_editor 是内容一致的可复制模板。", LIGHT_GOLD, GOLD)
    doc.add_heading("4.2 创建与编辑行为树", level=2)
    add_number(doc, "点击 New Tree，再点击 Add Root 或从 Entry 分组拖入 Root。每棵树只允许一个 Root。")
    add_number(doc, "从 Palette 将 Sequence、Selector、Random Selector、Parallel、Repeat、Action、Condition、Wait 或 Decorator 拖入画布，也可以右键画布创建。")
    add_number(doc, "从父节点底部方块拖向子节点顶部方块完成连接。一个 Sequence 或 Selector 可以连接多个子节点。")
    add_number(doc, "横向拖动同级节点调整执行顺序；Sequence 和 Selector 按从左到右执行。")
    add_number(doc, "选中节点，在右侧 Node Inspector 修改 Title、Type、Description 和类型化 Parameters；只有自定义字段需要展开 Advanced JSON。")
    add_number(doc, "右键节点可以断开、删除、启用/禁用或折叠；工具栏 Undo/Redo 用于恢复编辑操作。")
    add_number(doc, "填写以 res:// 开头的资源路径并点击 Save Tree。保存前会自动进行结构校验。")
    doc.add_heading("4.3 配置 Action 与 Condition", level=2)
    add_table(doc, ["目标", "节点参数示例", "Actor 端要求"], [
        ["执行巡逻", '{"action_name":"patrol"}', "实现 patrol(blackboard, delta, node)"],
        ["执行攻击", '{"action_name":"attack_right"}', "返回 SUCCESS / FAILURE / RUNNING"],
        ["方法条件", '{"condition_name":"can_attack"}', "方法可返回 bool 或 BTStatus"],
        ["黑板条件", "Key=has_target, Operator=Equals, Value=true", "Actor/Runner 先写入 has_target"],
    ], [1.05, 2.65, 2.8], 8.4)
    doc.add_heading("4.4 将行为树绑定到 NPC", level=2)
    add_number(doc, "在 NPC 场景下添加 BehaviorTreeComponent 节点。")
    add_number(doc, "把已保存的 BTTreeResource 拖入 behavior_tree 属性。")
    add_number(doc, "组件是 NPC 子节点时 actor_path 可留空；否则设置指向 Actor 的 NodePath，或在代码中调用 set_actor。")
    add_number(doc, "角色移动通常启用 tick_on_physics；纯逻辑更新可以使用 tick_on_process。避免同时开启导致双重 tick。")
    add_number(doc, "需要暂停、继续或重置时调用 stop_tree()、start_tree()、restart_tree()。更换 behavior_tree 会自动清理旧节点记忆。")

    doc.add_heading("4.5 Live Debug 使用", level=2)
    add_number(doc, "在 Behavior Tree 面板通过资源选择器打开 NPC 正在使用的同一个 .tres。")
    add_number(doc, "保持 Live Debug 勾选并运行游戏。Runner 会周期性发布原子 JSON 快照。")
    add_number(doc, "Active Path 显示当前链条；绿色、红色或运行色标识当前叶节点结果。")
    add_number(doc, "启用 Dim Inactive 淡化非当前分支；Failure Reasons 在节点卡片和 Failures 菜单显示失败原因。")
    add_number(doc, "启用 Blackboard 展开编辑器内四列表格，查看当前 Actor 的 Key、运行时类型、值和 Schema 状态；该面板不会出现在游戏画面。")
    add_number(doc, "Runtime Path 中的按钮可定位节点；即使节点位于折叠子树内，也会自动展开并聚焦。")
    add_picture(doc, VISUAL / "07_runtime_failure.png", 6.35, "图 4  Live Debug：当前失败链、非活动分支淡化和 Action 失败原因")
    add_picture(doc, VISUAL / "07b_live_blackboard.png", 6.35, "图 5  Live Blackboard：声明键、动态键与类型错误")


def add_demo(doc: Document) -> None:
    doc.add_heading("5 示例游戏与验证场景", level=1)
    add_body(doc, "测试项目提供可直接运行的 2D 场景 res://scenes/test_game.tscn。玩家支持移动、近战、冲刺、治疗和隐身；三个敌人均由 BehaviorTreeComponent 驱动。EnemyA 使用基础巡逻战斗树，EnemyB 与 EnemyC 使用复杂守卫验证树。")
    doc.add_heading("5.1 玩家操作", level=2)
    add_table(doc, ["按键", "功能", "验证目标"], [
        ["A/D 或方向键", "左右移动", "触发敌人的方向判断、追击和攻击"],
        ["J / 鼠标左键", "近战攻击", "伤害敌人并触发低血量恢复分支"],
        ["Space", "冲刺与短暂无敌", "验证连续运行 Action 和资源消耗"],
        ["H", "使用治疗次数", "验证玩家状态和 UI"],
        ["C", "隐身 2 秒", "打断索敌并触发最后已知位置搜索"],
        ["T", "暂停/继续全部 Runner", "验证 start/stop 生命周期"],
        ["R", "重置场景与行为树", "验证 restart 和记忆清理"],
    ], [1.1, 1.65, 3.75], 8.4, {0})
    doc.add_heading("5.2 复杂守卫的六个优先级分支", level=2)
    add_table(doc, ["优先级", "分支", "条件与行为"], [
        ["1", "Emergency Recovery", "低血量且检测到玩家时先撤退，再延迟治疗"],
        ["2", "Attack Left", "玩家在左侧攻击距离内；含 Condition、Cooldown、Time Limit"],
        ["3", "Attack Right", "玩家在右侧攻击距离内；含 Condition、Cooldown、Time Limit"],
        ["4", "Chase", "已索敌但不在攻击距离，主动向玩家前进"],
        ["5", "Search", "玩家隐身或丢失后移动到 last_known_player_x"],
        ["6", "Patrol", "未索敌时由 Random Selector 选择方向，Parallel 同时执行移动与 0.25 秒 Wait"],
    ], [0.65, 1.7, 4.15], 8.4, {0})
    doc.add_heading("5.3 推荐演示流程", level=2)
    add_number(doc, "远离敌人观察 Patrol，并在编辑器看到巡逻链。")
    add_number(doc, "进入 330 像素索敌范围，观察 Reactive Selector 从巡逻切换到 Chase。")
    add_number(doc, "靠近敌人左右两侧，观察方向攻击和 Cooldown 失败原因。")
    add_number(doc, "按 C 隐身，观察目标锁定解除并进入 Search Last Known。")
    add_number(doc, "将复杂敌人打到临界血量，观察 Retreat 和 Heal。")
    add_number(doc, "切换显示开关，对比路径高亮、淡化、小地图、搜索、鱼眼和折叠效果。")
    add_callout(doc, "目标检测参数", "敌人在 330 像素内获得可见玩家，超过 460 像素才释放目标。该滞回区间用于避免追击和巡逻在边界处频繁抖动。")


def add_testing(doc: Document) -> None:
    doc.add_heading("6 测试方法与结果", level=1)
    add_body(doc, "测试采用六层验证，并使用严格日志判定：非零退出码、ERROR、SCRIPT ERROR、FAIL、对象泄漏和原生崩溃回溯都视为失败。不能仅依据脚本最终退出码判断通过，因为 Godot 在部分脚本错误下仍可能返回 0。")
    doc.add_heading("6.1 自动测试总表", level=2)
    add_table(doc, ["测试层", "断言", "结果", "主要覆盖"], [
        ["资源与运行时", "153", "153/153", "节点语义、Decorator、Schema 引用、Actor、生命周期与调试桥"],
        ["编辑器 GUI", "185", "185/185", "图编辑、历史、Schema key picker、19 个开关、单线交互与 Live Blackboard"],
        ["基础游戏", "13", "13/13", "3 个敌人、攻击伤害、巡逻、死亡重生与无 Timer 泄漏"],
        ["复杂竞技场", "26", "26/26", "新节点、Schema、索敌、搜索、撤退、治疗和失败原因"],
        ["真实视觉回归", "70", "70/70", "17 张核心 1600×900 GPU 截图、灰度、单线、复杂树与 Schema key picker"],
        ["合计", "447", "447/447", "全部核心自动断言通过；另有性能 511/511、发布包 53/53"],
    ], [1.15, 0.7, 0.95, 3.7], 8.2, {1, 2})
    add_callout(doc, "稳定性复验", "最终日志中 FAIL、ERROR、SCRIPT ERROR、对象泄漏和原生崩溃均为 0；复杂场景保持 3 个活动 Runner，并通过 180 个物理帧自然运行烟雾测试。", LIGHT_TEAL, TEAL)
    doc.add_heading("6.2 覆盖的关键边界", level=2)
    add_bullet(doc, "重复 ID、空资源、错误 Root、悬空父节点、父节点循环、不可达分支和非法 Decorator 归属。")
    add_bullet(doc, "空 Root、空 Sequence、空 Selector、禁用节点、缺失 Actor 方法和所有黑板比较运算符。")
    add_bullet(doc, "Reactive Selector 抢占、RUNNING 索引恢复、Runner 启停重启、Process/Physics tick 和热切换资源。")
    add_bullet(doc, "创建/删除/连接/断开/拖拽的撤销重做，删除子树时附属 Decorator 的一致性。")
    add_bullet(doc, "截断 Live Debug JSON、多 Runner 合并发布、临时文件清理和完整快照恢复。")
    add_bullet(doc, "Blackboard Schema 默认值、五种类型、严格/动态键策略、深复制、保存加载、错误快照和编辑器表格。")
    add_bullet(doc, "鱼眼与滚轮缩放、关闭恢复、折叠与搜索、路径与失败标注、连线方式及组合开关残留。")
    doc.add_heading("6.3 视觉回归", level=2)
    add_two_pictures(doc, VISUAL / "02_compact.png", VISUAL / "05_collapsed.png",
                     "图 5a  Compact 模式", "图 5b  子树折叠和摘要")
    add_two_pictures(doc, VISUAL / "08_orthogonal_edges.png", VISUAL / "09_bundled_edges.png",
                     "图 5c  正交连线", "图 5d  连线捆绑")
    add_body(doc, "视觉测试使用 NVIDIA GeForce RTX 5070 Laptop GPU 的 OpenGL 3.3 Compatibility 渲染器，而不是 Godot 无头 Dummy Rendering。17 张核心截图检查固定尺寸、非空颜色分布、形状与灰度编码、无障碍配色、单一连接线、Schema 面板/key picker 和复杂树；另有 4 张固定人工实验场景截图通过 21/21 条检查。人工复验覆盖基线、鱼眼关闭恢复、key picker、复杂树、364 节点搜索和七节点 Live Debug 活动链。")


def add_optimization(doc: Document, display: dict[int, dict[str, dict[str, float]]],
                     branch_rows: list[dict[str, str]], minimap_rows: list[dict[str, str]]) -> None:
    doc.add_heading("7 显示优化实验与量化数据", level=1)
    add_body(doc, "自动基准生成 31、121 和 364 节点三档行为树。所有方法在相同树生成规则下测量可见节点、卡片总面积、信息字段、降噪节点、局部放大倍率和更新耗时。比例表示界面信息或面积的变化，不直接等同于人的可读性提升。")
    add_picture(doc, ASSETS / "optimization_ratios.png", 6.2, "图 6  364 节点树中主要显示优化的量化比例")
    large = display[364]
    baseline = large["Baseline"]
    compact_drop = (1 - large["Compact Cards"]["card_area_px2"] / baseline["card_area_px2"]) * 100
    focus_drop = (1 - large["Subtree Focus"]["visible_ratio"]) * 100
    collapse_drop = (1 - large["Subtree Collapse"]["visible_ratio"]) * 100
    fisheye_gain = (large["Fisheye"]["focus_scale_gain"] - 1) * 100
    metric_rows = [
        ["Compact Cards", "卡片总面积", f'{baseline["card_area_px2"]:,.0f} → {large["Compact Cards"]["card_area_px2"]:,.0f} px²', f"降低 {compact_drop:.1f}%", "保留全部 364 节点"],
        ["Semantic Zoom", "信息字段", "6 → 1", "降低 83.3%", "低缩放仅保留关键身份信息"],
        ["Search Highlight", "降噪节点", "363 / 364", "99.7%", "目标节点保持高亮"],
        ["Subtree Focus", "可见节点", f'364 → {int(large["Subtree Focus"]["visible_nodes"])}', f"降低 {focus_drop:.1f}%", f'重建约 {large["Subtree Focus"]["rebuild_ms"]:.1f} ms'],
        ["Subtree Collapse", "可见节点", f'364 → {int(large["Subtree Collapse"]["visible_nodes"])}', f"降低 {collapse_drop:.1f}%", f'重建约 {large["Subtree Collapse"]["rebuild_ms"]:.1f} ms'],
        ["Fisheye", "目标放大", f'1.00 → {large["Fisheye"]["focus_scale_gain"]:.2f}', f"增加 {fisheye_gain:.1f}%", "保持全树上下文"],
        ["Branch Dimming", "非活动分支", "361 / 364", "淡化 99.2%", "透明度 1.00 vs 0.24"],
        ["Enhanced Minimap", "节点覆盖", "364 / 364", "覆盖 100%", "固定 230×150"],
    ]
    add_table(doc, ["方法", "指标", "实测值", "比例", "解释"], metric_rows,
              [1.28, 1.05, 1.55, 1.0, 1.62], 7.8, {3})
    doc.add_heading("7.1 三档规模对比", level=2)
    scale_rows = []
    for size in (31, 121, 364):
        methods = display[size]
        scale_rows.append([
            str(size),
            f'{int(methods["Subtree Focus"]["visible_nodes"])} ({(1-methods["Subtree Focus"]["visible_ratio"])*100:.1f}%↓)',
            f'{int(methods["Subtree Collapse"]["visible_nodes"])} ({(1-methods["Subtree Collapse"]["visible_ratio"])*100:.1f}%↓)',
            f'{int(methods["Search Highlight"]["dimmed_nodes"])}',
            f'{methods["Fisheye"]["focus_scale_gain"]:.2f}×',
        ])
    add_table(doc, ["总节点", "Focus 可见", "Collapse 可见", "Search 降噪", "鱼眼倍率"],
              scale_rows, [0.85, 1.55, 1.65, 1.45, 1.0], 8.5, {0, 1, 2, 3, 4})
    doc.add_heading("7.2 分支淡化与小地图性能", level=2)
    add_picture(doc, ASSETS / "performance_scaling.png", 6.25, "图 7  显示功能的一次性更新成本和稳态刷新成本")
    perf_rows = []
    for branch, mini in zip(branch_rows, minimap_rows):
        perf_rows.append([
            branch["tree_size"],
            f'{float(branch["dimmed_ratio"])*100:.1f}%',
            f'{float(branch["first_update_ms"]):.3f} ms',
            f'{float(branch["steady_update_ms"]):.5f} ms',
            f'{float(mini["coverage_ratio"])*100:.0f}%',
            f'{float(mini["steady_status_update_ms"]):.5f} ms',
        ])
    add_table(doc, ["节点", "淡化覆盖", "淡化首次", "淡化稳态", "小地图覆盖", "状态稳态"],
              perf_rows, [0.65, 1.05, 1.2, 1.2, 1.15, 1.25], 8.1, {0, 1, 2, 3, 4, 5})
    branch_large = next(row for row in branch_rows if row["tree_size"] == "364")
    minimap_large = next(row for row in minimap_rows if row["tree_size"] == "364")
    add_body(doc, f'在 364 节点树上，非活动分支透明度从 1.00 降为 0.24，亮度对比约 4.17 倍；本轮首次应用 {float(branch_large["first_update_ms"]):.3f} ms，稳态刷新 {float(branch_large["steady_update_ms"]):.5f} ms。增强小地图覆盖全部节点，本轮开关周期 {float(minimap_large["toggle_cycle_ms"]):.4f} ms，稳态状态更新 {float(minimap_large["steady_status_update_ms"]):.5f} ms。')
    doc.add_heading("7.3 对数据的正确解释", level=2)
    add_bullet(doc, "Compact、Focus 和 Collapse 的百分比是面积或可见节点减少率；Semantic Zoom 的百分比是字段减少率。")
    add_bullet(doc, "Fisheye 和 Minimap 不通过减少节点改善阅读，前者强化局部，后者提供空间定位，因此不应以节点减少率评价。")
    add_bullet(doc, "Focus 与 Collapse 需要重建视图，适合用户主动触发；分支淡化和路径高亮的稳态刷新成本更低，适合 Live Debug 持续更新。")
    add_bullet(doc, "自动指标证明界面信息量和程序性能，但不能独立证明认知负担下降。论文阶段建议增加 8–15 人的节点定位与路径解释实验。")
    add_callout(doc, "建议综合评价权重", "目标定位 30%，面积或可见节点降幅 25%，信息完整性 20%，交互延迟 15%，学习成本 10%。", LIGHT_GOLD, GOLD)
    doc.add_heading("7.4 人工对比实验准备", level=2)
    add_body(doc, "已完成被试内实验协议、固定 121/364 节点树、确定性七节点 Live Debug 快照和独立数据工作簿。工作簿为 15 名参与者预生成 270 次试验，包含六种循环平衡方法顺序、三种任务区组顺序、数据验证、公式汇总和预先冻结的排除规则。")
    add_picture(doc, HUMAN_VISUAL / "04_tree_364_live_debug_path.png", 6.35, "图 8  人工实验固定 364 节点树与七节点 Live Debug 执行链")
    add_callout(doc, "结论边界", "参与者尚未招募。当前只能报告协议、材料、公式和视觉验证已经完成，不能声称任何显示方法已经改善人的可读性或认知负担。", LIGHT_GOLD, GOLD)


def add_quality_and_future(doc: Document) -> None:
    doc.add_heading("8 问题修复、局限与后续工作", level=1)
    doc.add_heading("8.1 已定位并修复的问题", level=2)
    add_table(doc, ["问题", "影响", "修复结果"], [
        ["Live Debug 直接覆盖 JSON", "编辑器可能读取半写入内容", "按进程临时文件、快照合并和原子发布；截断文件静默等待"],
        ["start_tree 覆盖显式 Actor", "Action 重启后找不到方法", "set_actor 保持优先级；仅 ActorPath 变更时重新解析"],
        ["切换树保留旧记忆", "新树从旧索引继续执行", "behavior_tree setter 清理节点记忆和调试状态"],
        ["禁用节点不在调试路径", "无法定位失败节点", "记录 FAILURE、失败原因、完整路径和节点信号"],
        ["非法 JSON 输出红色错误", "编辑参数时持续报错", "使用无噪声解析，保留旧参数并显示英文提示"],
        ["GraphNode scale 鱼眼不稳定", "滚轮缩放、中心偏移和关闭残留", "改为内部尺寸/字体放大，中心补偿并清理旧变换"],
        ["复杂树自动布局重叠", "多 Decorator 卡片压住下一层", "按每层最大实际高度动态安排 Y，并允许 Fit 缩放到 0.10"],
        ["低缩放搜索定位偏移", "Inspector 选中目标但卡片不在视口", "按 GraphEdit zoom 换算 scroll_offset，并增加 0.10 缩放远距目标回归"],
        ["长运行路径相互覆盖", "Runtime Path 与 Selection 在 1600×900 重叠", "拆成两条独立滚动行；顶栏省略并保留完整 tooltip"],
        ["右巡逻 Wait 参数覆盖", "0.25 秒被错误 action_name 覆盖", "删除重复字段，并测试左右 Wait 都为 0.25 秒"],
        ["角色异步 Timer 泄漏", "测试/场景提前退出留下 SceneTreeTimer", "受伤和重生改为节点本地状态机；verbose 复验无泄漏"],
    ], [1.65, 2.05, 2.8], 8.1)
    doc.add_heading("8.2 当前局限", level=2)
    add_bullet(doc, "未实现 Service 节点；这是当前毕业设计范围的主动取舍。")
    add_bullet(doc, "Blackboard 已提供类型化键选择、自由输入、引用位置与未使用键分析；尚未提供跨树全局重命名重构。")
    add_bullet(doc, "Live Debug 使用本地 JSON 桥，适合编辑器与本机游戏进程，不面向网络远程调试。")
    add_bullet(doc, "显示优化已有自动指标、视觉回归、实验协议和 270 行数据模板，但尚未招募正式参与者。")
    add_bullet(doc, "0.9.0 ZIP、MIT 许可证、哈希清单和洁净安装已经完成；Godot Asset Library 页面、正式图标和跨版本矩阵仍待准备。")
    doc.add_heading("8.3 后续工作建议", level=2)
    add_number(doc, "在现有键选择与引用分析基础上增加安全的跨树键重命名重构。")
    add_number(doc, "增加 Subtree 节点或子树资源引用，支持大型行为树模块复用。")
    add_number(doc, "增加断点、逐 tick、历史路径时间轴和运行快照录制，提高调试深度。")
    add_number(doc, "按现成协议招募 8–15 人，执行 Baseline、Compact、Collapse、Fisheye、Search 和 Minimap 的 270 次预生成试验并分析真实数据。")
    add_number(doc, "完善版本迁移、API 文档、Asset Library 页面、插件图标与 Godot 4.x 兼容性矩阵，准备公开发布。")

    doc.add_heading("9 结论", level=1)
    add_body(doc, "本项目已经完成一个可在 Godot 4.6 中实际使用的可视化行为树组件。它不仅提供节点图编辑界面，还包含资源模型、结构校验、Actor 方法调用、黑板判断、Decorator、运行时状态机、组件化 NPC 绑定和编辑器 Live Debug。基础游戏和复杂竞技场证明 NPC 的巡逻、索敌、追击、左右攻击、搜索最后位置、撤退和治疗均由行为树驱动。")
    add_body(doc, "显示优化部分形成了可独立启停的 19 项功能，并通过三档规模基准和真实 GPU 截图验证。364 节点树上的主要结果包括：Compact 卡片面积减少 55.9%，Semantic Zoom 信息字段减少 83.3%，Subtree Focus 可见节点减少 97.5%，Subtree Collapse 减少 99.5%，Search 对 99.7% 非目标节点降噪，Fisheye 局部放大 20%，增强小地图覆盖 100% 节点。新增形状/图标、色盲安全配色和单连接显示均有独立开关与恢复测试。")
    add_body(doc, "运行时拓扑缓存使用 31、121、364 节点和 1、10、50 个 Runner 进行 126 次观测。加入每 Tick 精确拓扑签名后，中位数收益分别为 15.98%–26.46%、52.61%–56.02% 和 76.65%–78.00%；511/511 条性能断言通过。精确校验降低部分小树收益，但能安全发现同数量换序、改父级、Decorator 归属和节点实例替换；结果仅代表当前机器和固定工作负载。")
    add_callout(doc, "最终判断", "当前程序已经达到毕业设计中“基本功能完整、能够真实运行、能够可视化调试、具有显示优化研究与量化验证”的目标。后续重点应转向正式用户实验、子树复用和公开发布准备。")


def add_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("附录 A 关键资源与文件", level=1)
    add_table(doc, ["资源", "路径/说明"], [
        ["插件", "res://addons/behavior_tree_editor/"],
        ["基础敌人树", "res://behavior_trees/enemy_patrol_combat.tres"],
        ["复杂守卫树", "res://behavior_trees/complex_guard_validation_tree.tres"],
        ["复杂守卫 Schema", "res://behavior_trees/complex_guard_blackboard_schema.tres"],
        ["测试场景", "res://scenes/test_game.tscn"],
        ["插件使用说明", "testgame/testgame/BEHAVIOR_TREE_PLUGIN_USAGE.md"],
        ["视觉原始截图", "testgame/testgame/test_results/visual/"],
        ["人工实验协议", "research/display_optimization/human_comparison_study_protocol.md"],
        ["人工实验工作簿", "research/display_optimization/behavior_tree_human_comparison_study.xlsx"],
        ["显示原始数据", "testgame/testgame/test_results/ 下三份显示基准 CSV"],
        ["运行时性能原始数据", "testgame/testgame/test_results/runtime_profile.csv"],
        ["无人值守证据工作簿", "research/display_optimization/behavior_tree_unattended_run_evidence.xlsx"],
        ["可分发插件包", "dist/behavior-tree-editor-0.9.0.zip（MIT + SHA-256 清单）"],
    ], [1.7, 4.8], 8.6)
    doc.add_heading("附录 B 自动测试命令", level=1)
    commands = [
        "run_behavior_tree_tests.gd  -  资源与运行时 153 项",
        "run_editor_view_tests.gd  -  编辑器 GUI 185 项",
        "run_game_integration_tests.gd  -  基础游戏 13 项",
        "run_complex_arena_tests.gd  -  复杂竞技场 26 项",
        "run_visual_regression_tests.gd  -  真实 GPU 视觉 70 项",
        "run_human_study_visual_tests.gd  -  固定实验场景 GPU 视觉 21 项（附加）",
        "run_runtime_profile.gd  -  运行时拓扑缓存与性能 511 项",
        "package_behavior_tree_plugin.ps1  -  打包、哈希与洁净安装 53 项",
        "run_arena_smoke_test.gd  -  180 物理帧烟雾测试",
        "run_display_benchmarks.gd  -  31/121/364 节点显示基准",
    ]
    for item in commands:
        add_bullet(doc, item)
    add_body(doc, "无头脚本的一般命令格式：Godot_v4.6-stable_win64_console.exe --headless --path testgame/testgame --script res://tests/<脚本名>。视觉测试必须去掉 --headless，并使用 --rendering-method gl_compatibility；打包运行 tools/package_behavior_tree_plugin.ps1。测试证据目录由 .gdignore 隔离，不会被 Godot 当成重复插件源码扫描。")
    doc.add_heading("附录 C 数据口径", level=1)
    add_table(doc, ["字段", "含义"], [
        ["visible_ratio", "显示优化后可见节点数 / 原始总节点数"],
        ["card_area_px2", "所有可见节点 custom_minimum_size 面积之和"],
        ["information_fields", "当前信息密度层保留的代表性字段数"],
        ["dimmed_nodes", "被搜索或运行路径功能视觉降噪的节点数"],
        ["focus_scale_gain", "鱼眼目标节点的内部放大系数；当前实测为 1.20"],
        ["rebuild_ms", "目标显示操作后的视图重建或测量耗时；受机器负载影响"],
    ], [1.65, 4.85], 8.6)
    add_callout(doc, "数据可追溯性", "本报告中的比例均由 test_results 下的原始 CSV 重新计算；测试截图来自当前插件真实渲染，不使用示意图替代结果证据。", LIGHT_TEAL, TEAL)


def build() -> None:
    display = read_display_data()
    branch_rows = read_csv(RESULTS / "branch_dimming_results.csv")
    minimap_rows = read_csv(RESULTS / "enhanced_minimap_results.csv")
    create_charts(display, branch_rows, minimap_rows)

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    add_overview(doc)
    add_features(doc)
    add_display_features(doc)
    add_usage(doc)
    add_demo(doc)
    add_testing(doc)
    add_optimization(doc, display, branch_rows, minimap_rows)
    add_quality_and_future(doc)
    add_appendix(doc)
    add_running_furniture(doc)

    properties = doc.core_properties
    properties.title = "Godot 4.6 可视化行为树插件完整报告"
    properties.subject = "毕业设计项目功能、使用、测试与显示优化数据"
    properties.author = "Behavior Tree Plugin Project"
    properties.keywords = "Godot, Behavior Tree, Visual Programming, Live Debug, Display Optimization"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    # Normalize package relationships after the image-heavy build. This preserves
    # document content while avoiding a Microsoft Word open-time layout stall.
    normalized = Document(OUTPUT)
    normalized.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
