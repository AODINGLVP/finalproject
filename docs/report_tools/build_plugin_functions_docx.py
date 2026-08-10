from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "BEHAVIOR_TREE_PLUGIN_FUNCTIONS_ZH.md"
OUTPUT = ROOT / "docs" / "BEHAVIOR_TREE_PLUGIN_FUNCTIONS_ZH.docx"

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FA"
TEXT = "243447"
MUTED = "66788A"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, latin="Calibri", east_asia="Microsoft YaHei", size=11, bold=None, color=TEXT):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    width = table_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        table_pr.append(width)
    width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_inline(paragraph, text, bold=False):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, latin="Consolas", east_asia="Microsoft YaHei", size=9.5, color=DEEP_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, bold=bold)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DEEP_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("Godot 4.6 可视化行为树插件")
    set_font(run, size=25, bold=True, color=DEEP_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("完整功能说明与使用参考")
    set_font(run, size=16, bold=False, color=BLUE)

    box = doc.add_table(rows=4, cols=2)
    box.style = "Table Grid"
    values = [
        ("项目性质", "毕业设计 / Godot 4.6 编辑器插件"),
        ("当前版本", "0.9.0（当前工作区版本）"),
        ("活动项目", "testgame/testgame"),
        ("主要创新", "大型行为树的可读性与导航优化"),
    ]
    for row, values_row in zip(box.rows, values):
        for index, value in enumerate(values_row):
            row.cells[index].text = ""
            add_inline(row.cells[index].paragraphs[0], value, bold=index == 0)
            set_cell_fill(row.cells[index], LIGHT_BLUE if index == 0 else "FFFFFF")
    set_table_geometry(box, [1800, 7560])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("说明：插件 UI 保持英文；本文档使用中文解释界面、节点语义、运行时、Live Debug 和显示优化。")
    set_font(run, size=10, color=MUTED)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False


def parse_markdown_tables(lines, index):
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        rows.append([cell.strip() for cell in raw.split("|")])
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def widths_for(headers):
    count = len(headers)
    lowered = [h.lower() for h in headers]
    if count == 2:
        return [2200, 7160]
    if count == 3:
        if headers[0] == "节点":
            return [1500, 1600, 6260]
        if "#" in headers[0]:
            return [750, 1900, 6710]
        return [1800, 1800, 5760]
    if count == 4:
        return [650, 2350, 1900, 4460]
    return [TABLE_WIDTH_DXA // count] * count


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.rows[row_index].cells[col_index]
            cell.text = ""
            add_inline(cell.paragraphs[0], value, bold=row_index == 0)
            if row_index == 0:
                set_cell_fill(cell, LIGHT_BLUE)
            elif row_index % 2 == 0:
                set_cell_fill(cell, LIGHTER_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    if run.font.name != "Consolas":
                        run.font.size = Pt(8.4 if len(rows[0]) >= 4 else 9.2)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths_for(rows[0]))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code(doc, code_lines):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F1F4F7")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(code_lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_font(run, latin="Consolas", east_asia="Microsoft YaHei", size=8.8, color="263238")
    set_table_geometry(table, [TABLE_WIDTH_DXA])


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)
    add_cover(doc)

    header = doc.sections[-1].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Godot Behavior Tree Editor · 功能说明")
    set_font(run, size=8.5, color=MUTED)
    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("当前工作区文档 · 2026-08-10")
    set_font(run, size=8.5, color=MUTED)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 1  # Cover already contains the Markdown H1.
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_markdown_tables(lines, index)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = min(3, len(heading.group(1)) - 1)
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline(paragraph, heading.group(2), bold=True)
            index += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, text)
            index += 1
            continue
        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            index += 1
            continue
        paragraph = doc.add_paragraph()
        add_inline(paragraph, stripped.replace("  ", " "))
        index += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
